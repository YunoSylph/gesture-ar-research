from __future__ import annotations

import json
import math
import zipfile
from collections import Counter
from datetime import date
from pathlib import Path
from typing import Iterable

import numpy as np
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "generated"
FIG_DIR = OUT_DIR / "figures"
SCREENSHOT_DIR = OUT_DIR / "report_assets"
DOCX_PATH = OUT_DIR / "Yuzhakov_MV_AR_Gesture_Thesis_Draft_2026-05-18.docx"

ACCENT = RGBColor(31, 78, 121)
MUTED = RGBColor(95, 95, 95)
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"
LIGHT_GREEN = "E2F0D9"
LIGHT_RED = "FCE4D6"


def read_json(path: str) -> dict:
    return json.loads((ROOT / path).read_text(encoding="utf-8"))


def read_manifest_counts(path: str) -> tuple[int, Counter[str]]:
    counter: Counter[str] = Counter()
    total = 0
    p = ROOT / path
    with p.open(encoding="utf-8") as fh:
        for line in fh:
            if not line.strip():
                continue
            total += 1
            row = json.loads(line)
            counter[row.get("target_label") or row.get("label")] += 1
    return total, counter


def metrics_from_report(path: str) -> dict:
    data = read_json(path)
    rec = data["recognition"]
    lat = data["latency"]
    return {
        "accuracy": rec["accuracy"],
        "balanced_accuracy": rec["balanced_accuracy"],
        "macro_f1": rec["macro_f1"],
        "weighted_f1": rec["weighted_f1"],
        "median_ms": lat["offline_latency_ms_median"],
        "p95_ms": lat["offline_latency_ms_p95"],
        "n": lat["num_samples"],
        "confusion_matrix": rec.get("confusion_matrix"),
        "per_class": rec.get("per_class", {}),
    }


def fmt(value: float, digits: int = 4) -> str:
    return f"{value:.{digits}f}"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 10) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 24 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "BFBFBF")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, h in enumerate(headers):
        set_cell_text(hdr.cells[idx], h, bold=True, size=10)
        set_cell_shading(hdr.cells[idx], LIGHT_BLUE)
        if widths:
            hdr.cells[idx].width = Cm(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=10)
            if widths:
                cells[idx].width = Cm(widths[idx])
    doc.add_paragraph()
    return table


def add_formula_table(doc: Document, rows: list[list[str]]) -> None:
    """Add a compact academic formula table with readable math-like notation."""

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    headers = ["№", "Формула", "Применение в проекте"]
    widths = [1.2, 7.4, 7.2]
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, h in enumerate(headers):
        set_cell_text(hdr.cells[idx], h, bold=True, size=9)
        set_cell_shading(hdr.cells[idx], LIGHT_GREEN)
        hdr.cells[idx].width = Cm(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=9)
            cells[idx].width = Cm(widths[idx])
            if idx == 1:
                for paragraph in cells[idx].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = "Cambria Math"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
                        run.italic = False
            elif idx == 0:
                cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def get_font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/timesbd.ttf" if bold else "C:/Windows/Fonts/times.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for item in candidates:
        if Path(item).exists():
            return ImageFont.truetype(item, size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = word if not current else current + " " + word
        bbox = draw.textbbox((0, 0), candidate, font=font)
        if bbox[2] - bbox[0] <= width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_box(draw, box, text, fill, outline="#4F81BD", font=None):
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=12, fill=fill, outline=outline, width=2)
    font = font or get_font(24, True)
    small = get_font(18)
    lines = wrap_text(draw, text, font, x2 - x1 - 30)
    if len(lines) > 2:
        font = small
        lines = wrap_text(draw, text, font, x2 - x1 - 30)
    line_h = font.size + 5 if hasattr(font, "size") else 22
    total_h = line_h * len(lines)
    y = y1 + (y2 - y1 - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, y), line, fill="#1F1F1F", font=font)
        y += line_h


def draw_arrow(draw, start, end, color="#4F81BD") -> None:
    draw.line([start, end], fill=color, width=4)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 14
    p1 = (end[0] - size * math.cos(angle - math.pi / 6), end[1] - size * math.sin(angle - math.pi / 6))
    p2 = (end[0] - size * math.cos(angle + math.pi / 6), end[1] - size * math.sin(angle + math.pi / 6))
    draw.polygon([end, p1, p2], fill=color)


def make_architecture_figure(path: Path) -> None:
    img = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(img)
    title = get_font(40, True)
    draw.text((60, 38), "Архитектура текущего исследовательского конвейера", fill="#1F4E79", font=title)
    boxes = [
        ((80, 150, 360, 270), "IPN Hand + локальные клипы", "#EAF3F8"),
        ((470, 150, 750, 270), "MediaPipe: 21x3 ключевые точки", "#E2F0D9"),
        ((860, 150, 1140, 270), "JSONL-описания + NPZ-фрагменты", "#FFF2CC"),
        ((1250, 150, 1530, 270), "C0 / C1 / C1-T", "#FCE4D6"),
        ((470, 410, 750, 530), "Эталонная проверка распознавания", "#EAF3F8"),
        ((860, 410, 1140, 530), "C2: контекстная политика", "#E2F0D9"),
        ((1250, 410, 1530, 530), "React + Three.js AR-интерфейс", "#FFF2CC"),
        ((860, 670, 1140, 790), "Переносимость ONNX / Core ML", "#F2F2F2"),
    ]
    for box, text, fill in boxes:
        draw_box(draw, box, text, fill)
    arrows = [
        ((360, 210), (470, 210)),
        ((750, 210), (860, 210)),
        ((1140, 210), (1250, 210)),
        ((1390, 270), (1390, 410)),
        ((1250, 470), (1140, 470)),
        ((860, 470), (750, 470)),
        ((1000, 530), (1000, 670)),
        ((1000, 270), (1000, 410)),
    ]
    for start, end in arrows:
        draw_arrow(draw, start, end)
    note_font = get_font(22)
    note = (
        "Смысл разделения: классификатор жестов оценивается отдельно, "
        "а затем проверяется как часть слоя взаимодействия с порогами, устойчивостью и паузой между командами."
    )
    for idx, line in enumerate(wrap_text(draw, note, note_font, 1380)):
        draw.text((100, 835 + idx * 28), line, fill="#595959", font=note_font)
    img.save(path)


def make_metric_chart(path: Path, metrics: dict[str, dict]) -> None:
    img = Image.new("RGB", (1600, 950), "white")
    draw = ImageDraw.Draw(img)
    title = get_font(40, True)
    label_font = get_font(22)
    small = get_font(18)
    draw.text((60, 40), "Сравнение вариантов распознавания на полном тестовом разбиении IPN", fill="#1F4E79", font=title)
    variants = [
        ("C0 правило", metrics["C0 full"], "#A6A6A6"),
        ("C1 RF", metrics["C1 RF full"], "#70AD47"),
        ("C1-T TCN", metrics["C1-T TCN full"], "#4472C4"),
        ("C1-T проверка", metrics.get("C1-T TCN validated", metrics["C1-T TCN full"]), "#8064A2"),
    ]
    chart_x, chart_y, chart_w, chart_h = 120, 160, 1000, 560
    draw.rectangle((chart_x, chart_y, chart_x + chart_w, chart_y + chart_h), outline="#D9D9D9", width=2)
    for i in range(0, 6):
        y = chart_y + chart_h - i * chart_h / 5
        draw.line((chart_x, y, chart_x + chart_w, y), fill="#EEEEEE", width=1)
        draw.text((70, y - 12), f"{i/5:.1f}", fill="#595959", font=small)
    group_w = chart_w / len(variants)
    bar_w = 80
    for i, (name, data, color) in enumerate(variants):
        x0 = chart_x + i * group_w + 80
        values = [("Точность", data["accuracy"]), ("Macro F1", data["macro_f1"]), ("Weighted F1", data["weighted_f1"])]
        for j, (_, value) in enumerate(values):
            h = value * chart_h
            x = x0 + j * (bar_w + 20)
            y = chart_y + chart_h - h
            shade = color
            if j == 1 and color == "#4472C4":
                shade = "#5B9BD5"
            if j == 2 and color == "#4472C4":
                shade = "#2F5597"
            draw.rectangle((x, y, x + bar_w, chart_y + chart_h), fill=shade)
            draw.text((x - 5, y - 28), f"{value:.3f}", fill="#1F1F1F", font=small)
        draw.text((x0, chart_y + chart_h + 25), name, fill="#1F1F1F", font=label_font)
    legend_y = 770
    for j, name in enumerate(["Точность", "Macro F1", "Weighted F1"]):
        x = 150 + j * 250
        draw.rectangle((x, legend_y, x + 28, legend_y + 28), fill=["#4472C4", "#5B9BD5", "#2F5597"][j])
        draw.text((x + 40, legend_y - 2), name, fill="#1F1F1F", font=label_font)
    latency_x = 1190
    draw.text((latency_x, 165), "p95 задержки, мс", fill="#1F4E79", font=get_font(28, True))
    for idx, (name, data, color) in enumerate(variants):
        y = 235 + idx * 95
        max_w = 300
        width = min(max_w, data["p95_ms"] / 60 * max_w)
        draw.text((latency_x, y), name, fill="#1F1F1F", font=label_font)
        draw.rectangle((latency_x, y + 34, latency_x + max_w, y + 58), outline="#D9D9D9")
        draw.rectangle((latency_x, y + 34, latency_x + width, y + 58), fill=color)
        draw.text((latency_x + max_w + 15, y + 28), f"{data['p95_ms']:.3f}", fill="#1F1F1F", font=label_font)
    draw.text(
        (latency_x, 560),
        "TCN не только точнее случайного леса\nпо macro F1, но и существенно быстрее\nпри пакетном выводе на полном тесте.",
        fill="#595959",
        font=label_font,
        spacing=6,
    )
    img.save(path)


def make_confusion_matrix(path: Path, matrix: list[list[int]], labels: list[str]) -> None:
    img = Image.new("RGB", (1500, 1200), "white")
    draw = ImageDraw.Draw(img)
    title = get_font(40, True)
    font = get_font(21)
    small = get_font(18)
    draw.text((70, 40), "Матрица ошибок C1-T TCN на полном тестовом разбиении IPN", fill="#1F4E79", font=title)
    m = np.asarray(matrix, dtype=float)
    max_val = max(1.0, float(m.max()))
    start_x, start_y = 300, 170
    cell = 105
    for i, label in enumerate(labels):
        draw.text((start_x + i * cell + 8, start_y - 55), label.replace("_", "\n"), fill="#1F1F1F", font=small, spacing=2)
        draw.text((70, start_y + i * cell + 36), label, fill="#1F1F1F", font=font)
    for r in range(len(labels)):
        for c in range(len(labels)):
            value = m[r, c]
            intensity = int(255 - 170 * (value / max_val))
            fill = (intensity, int(230 - 110 * value / max_val), 255)
            x1 = start_x + c * cell
            y1 = start_y + r * cell
            draw.rectangle((x1, y1, x1 + cell, y1 + cell), fill=fill, outline="#BFBFBF")
            text = str(int(value))
            bbox = draw.textbbox((0, 0), text, font=font)
            draw.text((x1 + (cell - bbox[2] + bbox[0]) / 2, y1 + 38), text, fill="#1F1F1F", font=font)
    draw.text((start_x, start_y + len(labels) * cell + 38), "Строки: истинный класс. Столбцы: предсказанный класс.", fill="#595959", font=font)
    img.save(path)


def make_live_card(path: Path, summary: dict) -> None:
    img = Image.new("RGB", (1500, 650), "white")
    draw = ImageDraw.Draw(img)
    title = get_font(38, True)
    head = get_font(30, True)
    body = get_font(24)
    draw.rounded_rectangle((50, 50, 1450, 600), radius=18, fill="#F7FBFE", outline="#9ECAE1", width=3)
    draw.text((90, 85), "Живой диагностический срез: веб-камера + ONNX", fill="#1F4E79", font=title)
    items = [
        ("Кадры", str(summary.get("frames", ""))),
        ("Длительность", f"{summary.get('duration_ms', 0)/1000:.2f} с"),
        ("FPS среднее / p95", f"{summary['fps']['mean']:.3f} / {summary['fps']['p95']:.3f}"),
        ("Обработка средняя / p95", f"{summary['processing_ms']['mean']:.3f} / {summary['processing_ms']['p95']:.3f} мс"),
        ("Метод", summary.get("method", "")),
        ("Источник", summary.get("source", "")),
    ]
    for idx, (k, v) in enumerate(items):
        x = 100 + (idx % 3) * 440
        y = 180 + (idx // 3) * 155
        draw.text((x, y), k, fill="#595959", font=body)
        draw.text((x, y + 42), v, fill="#1F1F1F", font=head)
    draw.text(
        (100, 515),
        "Пока живая сессия фиксирует техническую работоспособность потока; пользовательская валидация и локальная настройка C2 остаются следующими этапами.",
        fill="#595959",
        font=body,
    )
    img.save(path)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(0)

    for style_name, size, color in [
        ("Heading 1", 16, ACCENT),
        ("Heading 2", 15, ACCENT),
        ("Heading 3", 14, ACCENT),
    ]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def add_page_number(section) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def add_title_page(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    def center(text: str, size: int = 14, bold: bool = False, spacing_after: int = 0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(spacing_after)
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(size)
        r.bold = bold
        return p

    center("Федеральное государственное бюджетное образовательное учреждение высшего образования", 13)
    center("«Казанский национальный исследовательский технический университет им. А.Н. Туполева-КАИ»", 13)
    center("(КНИТУ-КАИ)", 13, spacing_after=6)
    center("Институт компьютерных технологий и защиты информации", 13)
    center("Кафедра систем информационной безопасности", 13)
    doc.add_paragraph()
    doc.add_paragraph()
    center("О Т Ч Е Т", 18, True, spacing_after=10)
    center("по результатам научно-исследовательской работы", 16, True)
    center("рабочая редакция магистерской диссертации", 15)
    doc.add_paragraph()
    center(
        "Тема: «Контекстно-зависимое распознавание жестов по временным последовательностям ключевых точек кисти для бесконтактного взаимодействия в дополненной реальности»",
        14,
        True,
    )
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(8)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for line in [
        "Выполнил: студент гр. 4267 Южаков М.В.",
        "Руководитель практики: Аникин И.В.",
        "Срез проекта: 18.05.2026",
        "Оценка ______________",
    ]:
        run = p.add_run(line + "\n")
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(14)

    center("Казань - 2026", 14)
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.first_line_indent = Cm(0)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_para(doc: Document, text: str, bold_start: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if bold_start and text.startswith(bold_start):
        r = p.add_run(bold_start)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        rest = p.add_run(text[len(bold_start) :])
        rest.font.name = "Times New Roman"
        rest.font.size = Pt(14)
        rest._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    else:
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run("• " + item)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(14)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(12)
    r.italic = True


def add_picture_if_exists(doc: Document, path: Path, width_inches: float, caption: str) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Inches(width_inches))
        add_caption(doc, caption)


def normalized_jpeg(src: Path, dst: Path) -> Path:
    if not src.exists():
        return src
    with Image.open(src) as img:
        img.convert("RGB").save(dst, "JPEG", quality=92, optimize=True)
    return dst


def add_literature(doc: Document) -> None:
    add_heading(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 1)
    refs = [
        "Koutsabasis P., Vogiatzidakis P. Empirical Research in Mid-Air Interaction: A Systematic Review // International Journal of Human-Computer Interaction. 2019. Vol. 35, No. 18. P. 1747-1768.",
        "Nguyen R., Gouin-Vallerand C., Amiri M. Hand interaction designs in mixed and augmented reality head mounted display: a scoping review and classification // Frontiers in Virtual Reality. 2023. Vol. 4. Article 1171230.",
        "Gavgiotaki D., Ntoa S., Margetis G., Apostolakis K.C., Stephanidis C. Gesture-based Interaction for AR Systems: A Short Review // PETRA '23. 2023. P. 284-292.",
        "Herbert O.M. et al. Static and Dynamic Hand Gestures: A Review of Techniques of Virtual Reality Manipulation // Sensors. 2024. Vol. 24, No. 12. Article 3760.",
        "Wang P., Wang S., Wang Y., Billinghurst M., Yang H. et al. Exploring the Role of Hand Gestures in AR/MR Remote Collaboration for Industry: A State-of-the-Art Review // Results in Engineering. 2025. Vol. 28. Article 107655.",
        "Piumsomboon T., Clark A., Billinghurst M., Cockburn A. User-Defined Gestures for Augmented Reality // INTERACT 2013. LNCS 8118. P. 282-299.",
        "Sluÿters A., Sellier Q., Nigay L. Consistent, Continuous, and Customizable Mid-Air Gesture Interaction for Browsing Multimedia Objects on Large Displays // International Journal of Human-Computer Interaction. 2023. Vol. 39, No. 1. P. 1-23.",
        "Ackad M., Schneider O., Isenberg O., Graham T.C.N., Klokmose C. An In-the-Wild Study of Learning Mid-Air Gestures to Browse a Hierarchical Information Space // CHI '15. 2015. P. 1329-1338.",
        "Buchmann V., Violich S., Billinghurst M., Cockburn A. FingARtips: Gesture Based Direct Manipulation in Augmented Reality // GRAPHITE '04. 2004. P. 212-221.",
        "Lee J.Y., Rhee G.W., Seo D.W. Hand Gesture-Based Tangible Interactions for Manipulating Virtual Objects in a Mixed Reality Environment // International Journal of Advanced Manufacturing Technology. 2010. Vol. 51. P. 1069-1082.",
        "Hürst W., van Wezel C. Gesture-Based Interaction via Finger Tracking for Mobile Augmented Reality // Multimedia Tools and Applications. 2013. Vol. 62. P. 233-258.",
        "Radkowski R., Herrema J., Oliver G. Interactive Hand-Gesture-Based Assembly for Augmented Reality // International Journal of Human-Computer Interaction. 2012. Vol. 28, No. 8. P. 1-15.",
        "Sun W. et al. Tangible and Mid-Air Interactions in Hand-Held Augmented Reality for Upper Limb Rehabilitation: An Evaluation of User Experience and Motor Performance // International Journal of Advanced Manufacturing Technology. 2024.",
        "Lystbæk M.N., Rosenberg P., Pfeuffer K., Grønbæk J.E., Gellersen H. Gaze-Hand Alignment: Combining Eye Gaze and Mid-Air Pointing for Interacting with Menus in Augmented Reality // Proceedings of the ACM on Human-Computer Interaction. 2022. Vol. 6, ETRA. Article 145.",
        "Pourmemar S., Ng R.M.Y., Levin D.I.W., Stuerzlinger W. Predicting Human Performance in Vertical Hierarchical Menu Selection in Immersive Augmented Reality Using Hand-Gesture and Head-Gaze // IEEE VR. 2022. P. 28-37.",
        "Van den Bogaert L., Geerts D. User-Defined Mid-Air Haptic Sensations for Interacting with an AR Menu Environment // EuroHaptics 2020. LNCS 12272. P. 25-32.",
        "Satriadi K.A., Ens B., Cordeil M., Czauderna T., Willett W., Jenny B. Augmented Reality Map Navigation with Freehand Gestures // IEEE VR. 2019. P. 593-603.",
        "Kostic Z., Dumas C., Pratt S., Beyer J. Exploring Mid-Air Hand Interaction in Data Visualization // IEEE Transactions on Visualization and Computer Graphics. 2024. Vol. 30, No. 9. P. 1-15.",
        "Hincapié-Ramos J.D., Guo X., Moghadasian P., Irani P. Consumed Endurance: A Metric to Quantify Arm Fatigue of Mid-Air Interactions // CHI '14. 2014. P. 1063-1072.",
        "Cheema N. et al. Predicting Mid-Air Interaction Movements and Fatigue Using Deep Reinforcement Learning // CHI '20. 2020. P. 1-13.",
        "Li X., Han F., Sun X., Liu Y., Li Y., Chen Y. Bracelet: Arms-Down Selection for Kinect Mid-Air Gesture // Behaviour & Information Technology. 2019. Vol. 38, No. 4. P. 401-409.",
        "Reynaert V., Rekik Y., Berthaut F., Grisoni L. The Effect of Hands Synchronicity on Users Perceived Arms Fatigue in Virtual Reality Environment // International Journal of Human-Computer Studies. 2023. Vol. 177. Article 103092.",
        "Schön D., Menges A., Otterness N. Assessing Rotational Mid-Air Interactions for Augmented Reality // CHI '23. 2023. P. 1-15.",
        "Żywanowski K., Łysakowski M., Nowicki M.R. et al. Vision-Based Hand Pose Estimation Methods for Augmented Reality in Industry: Crowdsourced Evaluation on HoloLens 2 // Computers in Industry. 2025. Vol. 171. Article 104328.",
        "Bertolasi J., Garcia-Hernandez N.V., Memeo M., Guarischi M., Gori M. Evaluation of HoloLens 2 for Hand Tracking and Kinematic Features Assessment // Virtual Worlds. 2025. Vol. 4, No. 3. Article 31.",
        "Benitez-Garcia G., Olivares-Mercado J., Sanchez-Perez G., Yanai K. IPN Hand: A Video Dataset and Benchmark for Real-Time Continuous Hand Gesture Recognition // ICPR 2020. 2021. P. 4340-4347.",
        "Zhang F., Bazarevsky V., Vakunov A., Tkachenka A., Sung G., Chang C.-L., Grundmann M. MediaPipe Hands: On-device Real-time Hand Tracking. arXiv:2006.10214. 2020.",
        "Google AI Edge. Hand landmarks detection guide: MediaPipe Hand Landmarker. Last updated 2026-04-21. URL: https://ai.google.dev/edge/mediapipe/solutions/vision/hand_landmarker",
        "Bai S., Kolter J.Z., Koltun V. An Empirical Evaluation of Generic Convolutional and Recurrent Networks for Sequence Modeling. 2018. URL: https://vladlen.info/papers/TCN.pdf",
        "Chang C.-L., Uboweja E., Sung G., Sokal K., Grundmann M., Bazarevsky V. On-device Real-time Hand Gesture Recognition // ICCV Workshop on Computer Vision for Augmented and Virtual Reality. 2021.",
        "Benitez-Garcia G., Olivares-Mercado J., Sanchez-Perez G., Takahashi H. IPN HandS: Efficient Annotation Tool and Dataset for Skeleton-Based Hand Gesture Recognition // Applied Sciences. 2025. Vol. 15, No. 11. Article 6321.",
        "Warchocki J., Vlasenko M., Eisma Y.B. GRLib: An Open-Source Hand Gesture Detection and Recognition Python Library. arXiv:2310.14919. 2023.",
        "Noparlik R., Zdunek R. Hand Gesture Recognition System with Finite State Machine for Remote Desktop Control // SPLITECH. 2023.",
        "ONNX Runtime. CUDA Execution Provider documentation. URL: https://onnxruntime.ai/docs/execution-providers/CUDA-ExecutionProvider.html",
        "Apple. Core ML Tools: PyTorch Conversion Workflow. URL: https://apple.github.io/coremltools/docs-guides/source/convert-pytorch-workflow.html",
        "PyTorch. Get Started: Installing on Windows. URL: https://pytorch.org/get-started/locally/",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(f"{i}. {ref}")
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)


def add_toc(doc: Document) -> None:
    add_heading(doc, "СОДЕРЖАНИЕ", 1)
    entries = [
        ("Введение", "3"),
        ("1. Анализ предметной области", "5"),
        ("1.1. Бесконтактное взаимодействие и классификация жестов", "5"),
        ("1.2. Распознавание жестов по ключевым точкам кисти", "7"),
        ("1.3. Актуализация литературы", "8"),
        ("1.4. Выводы", "9"),
        ("2. Метод и архитектура текущего проекта", "10"),
        ("2.1. Постановка задачи и словарь жестов", "10"),
        ("2.2. Данные, описания клипов и тензорное представление", "11"),
        ("2.3. Модели C0, C1, C1-T и слой C2", "14"),
        ("2.4. Программная архитектура", "16"),
        ("2.5. Механики AR-сценариев", "18"),
        ("2.6. Выводы", "20"),
        ("3. Экспериментальная часть, новизна и результаты", "21"),
        ("3.1. План экспериментов", "21"),
        ("3.2. Результаты распознавания", "23"),
        ("3.3. Слой взаимодействия и живой диагностический срез", "26"),
        ("3.4. Перенос на телефонный AR-домен", "29"),
        ("3.5. Научная новизна и ограничения", "31"),
        ("Заключение", "32"),
        ("Список использованных источников", "33"),
    ]
    for name, page in entries:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0)
        run = p.add_run(name)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(14)
        p.add_run(" " + "." * max(3, 78 - len(name)) + " ")
        p.runs[-1].font.name = "Times New Roman"
        p.runs[-1]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        p.runs[-1].font.size = Pt(14)
        r2 = p.add_run(page)
        r2.font.name = "Times New Roman"
        r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r2.font.size = Pt(14)
    doc.add_page_break()


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)

    labels = ["no_gesture", "point_2f", "click_2f", "swipe_left", "swipe_right", "zoom_in", "zoom_out"]
    train_total, train_counts = read_manifest_counts("data/interim/manifests/ipn_train_full_landmarks.jsonl")
    test_total, test_counts = read_manifest_counts("data/interim/manifests/ipn_test_full_landmarks.jsonl")
    initial_train_total, _ = read_manifest_counts("data/interim/manifests/ipn_train_initial_landmarks.jsonl")
    initial_test_total, _ = read_manifest_counts("data/interim/manifests/ipn_test_initial_landmarks.jsonl")
    local_phone_total, local_phone_counts = read_manifest_counts("data/interim/manifests/local_phone_plan.jsonl")
    metrics = {
        "C0 full": metrics_from_report("artifacts/reports/ipn_c0_full_recognition.json"),
        "C1 RF full": metrics_from_report("artifacts/reports/ipn_c1_rf_full_recognition.json"),
        "C1-T TCN full": metrics_from_report("artifacts/reports/ipn_c1t_tcn_full_recognition.json"),
        "C1-T TCN validated": metrics_from_report("artifacts/reports/ipn_c1t_tcn_full_validated_recognition.json"),
        "C0 initial": metrics_from_report("artifacts/reports/ipn_c0_initial_recognition.json"),
        "C1 RF initial": metrics_from_report("artifacts/reports/ipn_c1_rf_initial_recognition.json"),
        "C1-T prototype initial": metrics_from_report("artifacts/reports/ipn_c1t_initial_recognition.json"),
        "C1-T TCN initial": metrics_from_report("artifacts/reports/ipn_c1t_tcn_initial_recognition.json"),
    }
    live_summary = read_json("artifacts/reports/live_session_summary.json")
    task_report = read_json("artifacts/reports/live_task_report.json")
    risk_report = read_json("artifacts/reports/recognition_risk_analysis.json")
    domain_report = read_json("artifacts/reports/domain_readiness.json")
    mobile_contract = read_json("artifacts/mobile/gesture_mobile_bundle/portability_contract.json")
    preprocess_contract = read_json("artifacts/mobile/gesture_mobile_bundle/preprocessing_contract.json")
    task_session = task_report.get("session", {})
    task_items = task_report.get("tasks", {})
    placement_task = task_items.get("placement") or next(iter(task_items.values()), {})
    risk_full = risk_report["variants"]["c1t_tcn_full"]
    risk_validated = risk_report["variants"]["c1t_tcn_full_validated"]
    screenshot_main = normalized_jpeg(
        SCREENSHOT_DIR / "ui_01_main_object_control.png", FIG_DIR / "ui_01_main_object_control.jpg"
    )
    screenshot_sorting = normalized_jpeg(
        SCREENSHOT_DIR / "ui_02_virtual_sorting_task.png", FIG_DIR / "ui_02_virtual_sorting_task.jpg"
    )
    screenshot_advanced = normalized_jpeg(
        SCREENSHOT_DIR / "ui_03_advanced_controls.png", FIG_DIR / "ui_03_advanced_controls.jpg"
    )
    screenshot_results = normalized_jpeg(
        SCREENSHOT_DIR / "ui_04_results_page.png", FIG_DIR / "ui_04_results_page.jpg"
    )

    arch_path = FIG_DIR / "architecture.png"
    chart_path = FIG_DIR / "metrics.png"
    cm_path = FIG_DIR / "confusion_matrix.png"
    live_path = FIG_DIR / "live_summary.png"
    make_architecture_figure(arch_path)
    make_metric_chart(chart_path, metrics)
    make_confusion_matrix(cm_path, metrics["C1-T TCN validated"]["confusion_matrix"], labels)
    make_live_card(live_path, task_session)

    doc = Document()
    configure_styles(doc)
    add_title_page(doc)
    for section in doc.sections:
        add_page_number(section)

    add_toc(doc)

    add_heading(doc, "ВВЕДЕНИЕ", 1)
    add_para(
        doc,
        "Дополненная реальность постепенно переходит от демонстрационных приложений к практическим системам, в которых пользователь ожидает естественного, быстрого и устойчивого взаимодействия с виртуальными объектами. Одним из наиболее перспективных способов такого взаимодействия является бесконтактное управление жестами руки в воздухе: пользователь не держит контроллер, а выполняет команды движением кисти и пальцев перед камерой или гарнитурой дополненной реальности.",
    )
    add_para(
        doc,
        "Настоящий документ является расширенной рабочей редакцией отчета по результатам научно-исследовательской работы и фиксирует состояние магистерского проекта на 18.05.2026. По сравнению с предыдущей версией отчет дополнен текущими результатами разработки: подготовлен валидированный вариант временной нейросетевой модели, сформирован план переноса в телефонный AR-домен, создан мобильный пакет контрактов, добавлены отчеты о рисках распознавания и доменной готовности, а интерфейс теперь поддерживает набор типовых AR-сценариев, а не только демонстрационное управление одним объектом.",
    )
    add_para(
        doc,
        "Актуальность работы определяется двумя противоречиями. С одной стороны, обзоры по дополненной и смешанной реальности показывают устойчивый интерес к жестовому вводу и его применению в промышленности, навигации, визуализации и удаленном сотрудничестве [1]-[5]. С другой стороны, практическое качество такой системы зависит не только от классификатора, но и от всего конвейера: датасета, трекинга руки, временной модели, логики подавления ложных срабатываний и переносимости вывода модели на целевые платформы [24], [26]-[35].",
    )
    add_para(
        doc,
        "Цель работы - разработать и экспериментально оценить исследовательский конвейер, в котором основная подготовка и проверка выполняются на Windows, а последующий перенос рассчитан на мобильную дополненную реальность. Конвейер предназначен для контекстно-зависимого распознавания жестов по временным последовательностям ключевых точек кисти и для управления виртуальными объектами, меню, карточками, списками и измерительными сценами.",
    )
    add_para(doc, "Для достижения цели решаются следующие задачи:")
    add_bullets(
        doc,
        [
            "сформировать целевой словарь жестов, связанный с типовыми AR-действиями;",
            "подготовить публичный эталонный набор на основе IPN Hand и привести данные к единому формату: файл-описание JSONL и тензорные фрагменты NPZ;",
            "реализовать и сравнить несколько вариантов распознавания: правило-ориентированный базовый метод C0, классический машинный метод C1 и временную модель C1-T;",
            "выделить контекстно-зависимый слой C2, отвечающий за пороги уверенности, устойчивость распознавания, задержку между повторными действиями и подавление случайных команд;",
            "оценить метрики распознавания, показатели живых сессий и первые отчеты уровня пользовательских задач по AR-сценариям;",
            "подготовить переносимый контур ONNX/Core ML и мобильный пакет контрактов для будущей интеграции с телефонной AR-демонстрацией.",
        ],
    )
    add_para(
        doc,
        "Объект исследования - жестовое бесконтактное взаимодействие в системах дополненной реальности. Предмет исследования - методы распознавания и контекстной фильтрации жестов руки по временным последовательностям из 21 ключевой точки кисти с тремя координатами, получаемых из RGB-видео.",
    )
    add_para(
        doc,
        "Практическая значимость текущего этапа состоит в том, что создан не изолированный классификатор, а сквозная платформа для повторяемых экспериментов: от публичных данных до локального сервера распознавания, интерфейса дополненной реальности, журналов сессий, отчетов по сценариям и мобильного пакета контрактов. Поэтому результаты можно использовать как основу для финальной магистерской диссертации, дальнейшей локальной адаптации и пользовательского исследования.",
    )

    add_heading(doc, "1. АНАЛИЗ ПРЕДМЕТНОЙ ОБЛАСТИ", 1)
    add_heading(doc, "1.1. Бесконтактное взаимодействие и классификация жестов", 2)
    add_para(
        doc,
        "Жестовое бесконтактное взаимодействие в AR/MR рассматривается как форма ввода, при которой пользователь управляет цифровым содержимым движениями руки без физического контакта с поверхностью ввода. Для AR это особенно важно: использование мыши, клавиатуры или внешнего контроллера часто разрушает естественность сцены, а в производственных, медицинских или учебных сценариях может быть неудобно или небезопасно.",
    )
    add_para(
        doc,
        "В литературе выделяются статические и динамические жесты, манипуляционные и командные жесты, прямое и непрямое взаимодействие. Статические жесты проще стабилизировать, но они могут требовать удержания кисти в фиксированной позе. Динамические жесты лучше соответствуют действиям перелистывания, навигации и трансформации объектов, но требуют учета временной структуры движения [3], [4], [7].",
    )
    add_para(
        doc,
        "Для текущего проекта особенно важны три сценария: управление виртуальным объектом, навигация по элементам интерфейса и просмотр информационного содержимого. Эти сценарии покрывают типовые операции AR-прототипа: указание, подтверждение выбора, переход вперед/назад, масштабирование и подавление действия при отсутствии намеренного жеста.",
    )
    add_para(
        doc,
        "Отдельная проблема бесконтактного ввода - отсутствие тактильной опоры. Пользователь не получает естественного физического подтверждения, поэтому возрастает роль визуальной обратной связи: курсора, подсветки активной зоны, анимации подтверждения, индикаторов уверенности и статуса распознавания. Работы по меню в AR, визуализации и взаимодействию с картами показывают, что жесты должны оцениваться не только по точности классификации, но и по времени задачи, ошибкам, запоминаемости, утомляемости и субъективному удобству [14]-[18].",
    )
    add_para(
        doc,
        "Фактор утомляемости принципиален для магистерской постановки. Эффект «горилла-руки» возникает, когда пользователь долго держит руку в воздухе без опоры; это снижает точность, увеличивает число ошибок и ухудшает пользовательский опыт. Поэтому финальная оценка не должна ограничиваться общей точностью: необходимо фиксировать ложные срабатывания, повторные корректировки, задержку между жестом и действием, а также воспринимаемую физическую нагрузку [19]-[23].",
    )
    add_table(
        doc,
        ["Группа", "Примеры метрик", "Назначение в текущем проекте"],
        [
            ["Распознавание", "общая точность, сбалансированная точность, macro F1, weighted F1, матрица ошибок, задержка", "Сравнение C0, C1, C1-T на публичном эталоне"],
            ["Взаимодействие", "успешность задачи, частота ложных запусков, непреднамеренные действия, точность/полнота действий", "Оценка полезности C2 поверх классификатора"],
            ["Живая телеметрия", "FPS, время обработки, доля обнаружений, счетчики жестов и действий", "Контроль пригодности конвейера для камеры и AR-режима"],
            ["UX/эргономика", "SUS, UEQ-S, NASA-TLX или краткая шкала нагрузки", "Следующий этап пользовательского исследования"],
        ],
        widths=[3.0, 5.8, 7.2],
    )
    add_para(
        doc,
        "Для дальнейшего текста важно развести несколько близких терминов. Жест - это наблюдаемое движение или поза руки. Команда - это смысл, который система пытается извлечь из жеста. Действие - это уже результат в интерфейсе: переход к следующему объекту, подтверждение выбора, масштабирование или наведение курсора. Ошибка на уровне распознавания не всегда превращается в ошибку интерфейса, если контекстный слой успел ее подавить; и наоборот, формально правильный класс может быть неуместным, если он пришел слишком рано или повторился без намерения пользователя.",
    )
    add_table(
        doc,
        ["Термин", "Русское пояснение", "Как используется в проекте"],
        [
            ["Ключевые точки кисти", "21 нормализованная точка руки с координатами x, y, z", "служат входом для признаков позы и движения"],
            ["Окно последовательности", "набор кадров фиксированной длины T=32", "позволяет учитывать динамику жеста, а не только одну позу"],
            ["Маска кадров", "признак того, что рука найдена на конкретном кадре", "отделяет реальные точки от пустых кадров и влияет на live-диагностику"],
            ["Уверенность", "численная оценка надежности найденной руки или предсказанного класса", "используется в C2 для отсечения слабых команд"],
            ["Контекстный автомат", "логика состояний между распознаванием и действием", "требует устойчивости, применяет паузу и сброс по отсутствию команды"],
        ],
        widths=[4.1, 5.5, 6.2],
    )

    add_heading(doc, "1.2. Распознавание жестов по ключевым точкам кисти", 2)
    add_para(
        doc,
        "Возможны два крупных подхода к распознаванию жестов: работа с исходным RGB-видео и оптическим потоком либо работа с восстановленной позой руки. Подход по исходному изображению потенциально содержит больше визуальной информации, но дороже, сильнее зависит от сцены и сложнее переносится на мобильные устройства. Подход по ключевым точкам использует компактное представление кисти и позволяет строить более легкие модели, пригодные для частых экспериментов и вывода в реальном времени.",
    )
    add_para(
        doc,
        "MediaPipe Hands и современная задача MediaPipe Hand Landmarker формируют практическую основу такого подхода. В официальной документации указано, что Hand Landmarker работает со статическими изображениями, видео и потоком с камеры, возвращает ключевые точки в координатах изображения или мира и оценку принадлежности к левой/правой руке, а модель определяет 21 ключевую координату кисти [27], [28]. Для проекта это означает, что вход классификатора можно стабилизировать как тензор [T,21,3] с маской валидных кадров и численной уверенностью.",
    )
    add_para(
        doc,
        "Временная составляющая является ключевой для жестов перелистывания и масштабирования. Для таких жестов не хватает статической позы: важно направление, скорость, изменение расстояний между пальцами и устойчивость последовательности. Поэтому текущий проект разделяет признаки позы и признаки движения, а основной вариант C1-T реализован как компактная временная сверточная сеть. Выбор TCN обоснован тем, что сверточные последовательные модели могут служить сильной альтернативой рекуррентным архитектурам и проще оптимизируются на фиксированных временных окнах [29].",
    )
    add_para(
        doc,
        "Публичный датасет IPN Hand подходит для выбранной постановки, поскольку был создан для непрерывного распознавания жестов руки, содержит реальные сцены, широкий разброс условий и отдельный класс отсутствия жеста [26]. Новая работа IPN HandS 2025 дополнительно подтверждает значимость скелетной разметки и качества ключевых точек для обобщающей способности моделей [31]. Поэтому ориентация текущего проекта на первичное использование публичных данных является методологически оправданной.",
    )

    add_heading(doc, "1.3. Актуализация литературы", 2)
    add_para(
        doc,
        "Список источников из предыдущего отчета в целом остается релевантным для анализа предметной области: он покрывает бесконтактное взаимодействие, управление рукой в AR/MR, меню, визуализацию, усталость и трекинг руки. Однако текущий проект уже перешел от концепции к обучаемому конвейеру, поэтому литературу необходимо дополнить источниками по публичным датасетам, распознаванию по ключевым точкам, временным моделям и переносимому выводу модели.",
    )
    add_table(
        doc,
        ["Блок литературы", "Статус", "Решение"],
        [
            ["Обзоры бесконтактного взаимодействия и управления рукой в AR/MR [1]-[5]", "Актуальны", "Оставить как основу главы 1 и постановки проблемы"],
            ["Ранние работы по пользовательским жестам и прямой манипуляции [6], [9], [10]", "Актуальны как базовые", "Использовать для исторического и методического контекста"],
            ["Работы по меню, визуализации, обратной связи и утомляемости [14]-[23]", "Актуальны", "Сохранить для обоснования метрик уровня взаимодействия"],
            ["Источники по HoloLens/hand tracking 2025 [24], [25]", "Особенно актуальны", "Использовать как подтверждение современных ограничений трекинга"],
            ["Датасеты и извлечение ключевых точек", "Недостаточно в старом списке", "Добавить IPN Hand, IPN HandS, MediaPipe Hands/Hand Landmarker [26]-[28], [31]"],
            ["Временные модели и переносимость", "Недостаточно в старом списке", "Добавить TCN, ONNX Runtime, Core ML Tools, PyTorch [29], [34]-[36]"],
        ],
        widths=[5.2, 3.4, 7.2],
    )
    add_para(
        doc,
        "Таким образом, исходная библиография не требует радикальной замены, но требует расширения. В новом списке источников старые работы используются для описания интерфейсной и эргономической проблематики, а новые - для обоснования фактически реализованной исследовательской системы.",
    )

    add_heading(doc, "1.4. Выводы", 2)
    add_para(
        doc,
        "Анализ предметной области показывает, что исследовательская ценность работы находится не в простом факте распознавания жестов, а в проверке полной связки: извлечение ключевых точек, временное распознавание, контекстная политика действий и AR-интерфейс. Поэтому дальнейшие главы рассматривают систему как экспериментальную платформу, где классификатор и слой взаимодействия оцениваются раздельно и совместно.",
    )

    add_heading(doc, "2. МЕТОД И АРХИТЕКТУРА ТЕКУЩЕГО ПРОЕКТА", 1)
    add_heading(doc, "2.1. Постановка задачи и словарь жестов", 2)
    add_para(
        doc,
        "Текущий проект реализован как исследовательский конвейер, где первичная разработка ведется на Windows, а переносимость обеспечивается через ONNX/Core ML. Такое ограничение выбрано прагматически: основная подготовка данных, обучение, эталонные проверки и настольная демонстрация выполняются на Windows-ноутбуке с NVIDIA GPU, а мобильный/iOS-контур остается отдельным этапом переносимости.",
    )
    add_para(
        doc,
        "В целевой словарь включено семь классов: no_gesture, point_2f, click_2f, swipe_left, swipe_right, zoom_in, zoom_out. Они были выбраны как минимальный, но функционально полный набор для AR-сценариев: отсутствие команды, наведение, подтверждение, переход влево/вправо и масштабирование.",
    )
    add_table(
        doc,
        ["Класс", "Семантика", "IPN Hand mapping", "AR-действие"],
        [
            ["no_gesture", "отсутствие команды", "No gesture / D0X", "suppress/reset"],
            ["point_2f", "указание двумя пальцами", "Point-2f / B0B", "pointer/hover"],
            ["click_2f", "выбор/подтверждение", "Click-2f / G02", "select/confirm"],
            ["swipe_left", "движение влево", "Th-left / G05", "previous/rotate left"],
            ["swipe_right", "движение вправо", "Th-right / G06", "next/rotate right"],
            ["zoom_in", "масштабирование внутрь", "Zoom-in / G10", "zoom in"],
            ["zoom_out", "масштабирование наружу", "Zoom-o / G11", "zoom out"],
        ],
        widths=[3.0, 4.3, 4.0, 4.7],
    )
    add_para(
        doc,
        "Наличие класса no_gesture принципиально: в реальной AR-сцене пользователь постоянно двигает рукой, но не каждое движение является командой. Поэтому no_gesture используется не как «пустой» класс, а как reset/decay signal для context-aware политики.",
    )

    add_heading(doc, "2.2. Данные, описания клипов и тензорное представление", 2)
    add_para(
        doc,
        "Данные проекта организованы по контракту: JSONL-файл с описаниями клипов и NPZ-фрагменты с тензорами. Файл описания хранит метаданные клипа, метку, разбиение и путь к тензору; NPZ-фрагмент содержит ключевые точки [T,21,3], маску последовательности, уверенность по кадрам, оценку принадлежности к правой/левой руке и пространство координат. Такое разделение делает конвейер проверяемым: можно валидировать описания отдельно, переиспользовать извлеченные ключевые точки и запускать эталонные проверки без повторного чтения исходных видео.",
    )
    add_para(
        doc,
        f"На текущем этапе подготовлены три уровня данных. Первичный сбалансированный поднабор содержит {initial_train_total} обучающих и {initial_test_total} тестовых клипов, по 25/10 экземпляров на класс. Полный публичный эталонный набор содержит {train_total} обучающих и {test_total} тестовых клипа из целевых классов IPN Hand. Дополнительно создан план локального телефонного домена: {local_phone_total} будущих роликов, то есть по пять повторов на каждый из семи классов.",
    )
    rows = []
    for label in labels:
        rows.append([label, str(train_counts.get(label, 0)), str(test_counts.get(label, 0)), str(train_counts.get(label, 0) + test_counts.get(label, 0))])
    add_table(doc, ["Класс", "Train", "Test", "Всего"], rows, widths=[4.0, 3.0, 3.0, 3.0])
    add_para(
        doc,
        "Для обучения и вывода модели последовательности приводятся к длине T=32. Это значение является компромиссом между полнотой временного жеста и задержкой обработки. В дальнейших экспериментах целесообразно проверить окна 16/24/32, чтобы явно оценить компромисс между задержкой и точностью.",
    )
    add_para(
        doc,
        "Формально входной поток проекта можно описать как последовательность матриц ключевых точек. Такая запись полезна, потому что связывает видеопоток, извлечение ключевых точек и вход нейронной модели в одну воспроизводимую схему.",
    )
    add_formula_table(
        doc,
        [
            ["(1)", "P(t) = [p(t,1), ..., p(t,21)],  p(t,j) = (x(t,j), y(t,j), z(t,j));  X ∈ R^(T×21×3)", "Кадр t представлен 21 ключевой точкой кисти; вся запись является временным тензором жеста."],
            ["(2)", "s(t) = max(||p5_xy(t)-p17_xy(t)||_2, ||p9_xy(t)-p0_xy(t)||_2, ε)", "Масштаб ладони берется как максимум ширины и длины кисти; ε защищает от деления на ноль."],
            ["(3)", "p_norm(t,j) = (p(t,j) - p(t,0)) / s(t)", "Переход к координатам относительно запястья делает позу менее зависимой от положения руки в кадре."],
            ["(4)", "z(t) = [vec(P_norm(t)), c(t), w(t), Δc(t), Δw(t), s(t), Δs(t), q(t)] ∈ R^74;  Z ∈ R^(32×74)", "Вектор кадра объединяет 63 позных признака и 11 признаков глобального движения; именно Z подается в C1-T."],
        ],
    )
    local_rows = []
    for label in labels:
        local_rows.append([label, str(local_phone_counts.get(label, 0)), "phone_rear_ar", "ожидает записи видео"])
    add_table(
        doc,
        ["Класс", "План локальных роликов", "Домен", "Статус"],
        local_rows,
        widths=[3.4, 3.8, 3.2, 5.0],
    )
    add_para(
        doc,
        "Локальный телефонный набор не смешивается с IPN Hand автоматически. Он играет роль целевого домена: сначала публичная модель проверяется на IPN, затем те же веса тестируются на роликах с задней камеры телефона без дообучения, и только после этого допустимы калибровка порогов C2 или тонкая настройка модели. Такая последовательность нужна, чтобы не спутать улучшение от дополнительного обучения с реальной переносимостью распознавания.",
    )
    doc.add_picture(str(arch_path), width=Inches(6.4))
    add_caption(doc, "Рисунок 2.1 - Архитектура текущего исследовательского конвейера")

    add_heading(doc, "2.3. Модели C0, C1, C1-T и слой C2", 2)
    add_para(
        doc,
        "В проекте намеренно сохранены несколько уровней сложности. C0 является правило-ориентированной базовой моделью и нужен для нижней границы качества. C1 использует случайный лес на инженерных сводных признаках клипа и показывает, насколько далеко можно продвинуться без нейросетевой временной модели. C1-T использует компактную TCN-модель поверх временной последовательности признаков и является основным кандидатом для финальной системы.",
    )
    add_para(
        doc,
        "C1-T построен как стек временных блоков с одномерными свертками, нормализацией, нелинейностью GELU, прореживанием и остаточными связями. Конфигурация текущей модели: входная размерность 74, число классов 7, каналы (64, 64, 96), размер ядра 3, коэффициент прореживания 0.15. Такой размер модели хорошо соответствует цели проекта: не максимальная тяжеловесная видеосеть, а быстрый классификатор по ключевым точкам для многократных проверок гипотез и вывода в реальном времени.",
    )
    add_para(
        doc,
        "C2 не является еще одним классификатором. Это политика взаимодействия поверх предсказаний, реализованная как конечный автомат с порогом уверенности 0.62, требованием двух устойчивых кадров, паузой 250 мс между повторными действиями и сбросом после трех кадров без команды. Политика пропускает команду только после подтверждения устойчивости, подавляет низкую уверенность и не допускает слишком частых повторных действий.",
    )
    add_para(
        doc,
        "Для академического описания важно разделить три уровня: преобразование временного окна моделью, обучение через функцию потерь и контекстное превращение предсказания в действие интерфейса.",
    )
    add_formula_table(
        doc,
        [
            ["(5)", "H0 = Z^T;  H(l+1) = GELU(BN(Conv1D(k=3,d=2^l)(H(l)))) + Skip(H(l))", "Упрощенная запись временного блока TCN: свертка видит соседние кадры, а расширение d увеличивает временной охват."],
            ["(6)", "p_k = softmax(W·GAP(HL)+b)_k;  y_hat = argmax_k p_k", "После временного кодировщика выполняется усреднение по времени и выбор наиболее вероятного класса жеста."],
            ["(7)", "L = -(1/B) · Σ_b log p(y_b | Z_b)", "Кросс-энтропия задает критерий обучения классификатора на размеченных клипах."],
            ["(8)", "a(t)=g(y_hat(t)), если p_max(t)≥τ, stable(y_hat)≥S и t-t_last≥Δ; иначе a(t)=∅", "Формула C2: действие возникает только при достаточной уверенности, устойчивости и выдержанной паузе между командами."],
            ["(9)", "reset(t)=1, если count(no_gesture или p_max(t)<τ) ≥ R", "Сброс автомата по отсутствию жеста предотвращает зависание старой команды в интерфейсе."],
        ],
    )
    doc.add_page_break()
    add_table(
        doc,
        ["Вариант", "Назначение", "Преимущества", "Ограничения"],
        [
            ["C0 rule", "Нижняя граница качества", "Быстро, объяснимо, без обучения", "Слабая точность и переносимость правил"],
            ["C1 RF", "Классический машинный метод", "Высокая точность на сводных признаках", "Больше задержка, слабее учет времени"],
            ["C1-T TCN", "Основная временная модель", "Высокая macro F1 и низкая задержка", "Нужна проверка на локальном домене"],
            ["C2", "Слой «жест -> действие»", "Снижает случайные действия", "Нужна пользовательская оценка на уровне задач"],
        ],
        widths=[2.6, 4.0, 4.8, 4.6],
    )

    add_heading(doc, "2.4. Программная архитектура", 2)
    add_para(
        doc,
        "Репозиторий разделен на исследовательские модули данных, схем, предварительной обработки, моделей, оценки, взаимодействия, командной строки и серверного режима. Такое устройство позволяет запускать отдельные этапы через CLI и сохранять воспроизводимые артефакты в каталоге artifacts.",
    )
    add_para(
        doc,
        "Сервер живого режима реализован на FastAPI и WebSocket. Он поддерживает методы c0, c1_rf, c1t_tcn и onnx, может работать с воспроизведением сохраненного описания клипов или с потоком веб-камеры. В режиме камеры OpenCV получает кадры, MediaPipe HandLandmarker извлекает ключевые точки кисти, окно из 32 кадров превращается в тензор ключевых точек, после чего предсказатель возвращает метку, уверенность и распределение оценок по классам.",
    )
    add_para(
        doc,
        "Пользовательский интерфейс находится в demo/ar_interaction_app и реализован на React + Three.js. Текущая UI-поверхность содержит выбор метода распознавания, режимы воспроизведения датасета и потока камеры, прямое управление и режим C2, AR-сцену поверх камеры, панель результатов и live-телеметрию. Таким образом, проект уже позволяет проверять не только распознавание на сохраненных клипах, но и поведение в пользовательском контуре.",
    )
    add_picture_if_exists(doc, screenshot_main, 6.4, "Рисунок 2.2 - Главный экран Gesture AR в сценарии управления объектом")
    add_para(
        doc,
        "Важное изменение интерфейса состоит в том, что основной экран теперь не перегружен исследовательскими настройками. Пользователь сначала выбирает задачу, видит текущий сценарий, статус камеры, текущий жест, действие и телеметрию. Тонкие настройки - метод распознавания, источник данных, режим C2, параметры камеры и тестовые кнопки жестов - вынесены в раскрываемый блок. Это снижает когнитивную нагрузку и делает интерфейс ближе к экспериментальному стенду, где участник выполняет задачу, а не настраивает модель.",
    )
    add_picture_if_exists(doc, screenshot_advanced, 6.4, "Рисунок 2.3 - Расширенные настройки и live-статус интерфейса")
    add_para(
        doc,
        "Для переносимости подготовлен ONNX-экспорт модели ipn_c1t_tcn_full.onnx. На Windows также фиксируется Core ML contract stage: фактическая конвертация в mlpackage должна выполняться на поддерживаемой macOS/Linux-среде, но входные контракты и модельная архитектура уже отделены от Windows-only этапов.",
    )

    add_heading(doc, "2.5. Механики AR-сценариев", 2)
    add_para(
        doc,
        "Текущий интерфейс поддерживает тринадцать сценариев. В отчете важно описывать их не как декоративные режимы, а как разные классы пользовательских задач. Сценарий задает, какие действия считаются ожидаемыми, в каком порядке они должны появиться, какие жесты вызывают эти действия и какие признаки можно затем извлечь из журнала сессии. Поэтому сценарии связывают распознавание жестов с экспериментальной оценкой поведения.",
    )
    add_table(
        doc,
        ["Сценарий", "Пользовательская механика", "Типовая последовательность"],
        [
            ["Object Control", "наведение, выбор объекта, увеличение и уменьшение масштаба", "point -> click -> zoom in -> zoom out"],
            ["Gallery Navigation", "перелистывание виртуальных элементов и подтверждение выбранного", "right -> right -> left -> click"],
            ["AR Scroll List", "скроллинг плавающего списка и открытие строки", "right -> right -> left -> click"],
            ["Spatial Browser", "браузинг карточек, открытие центральной карточки и приближение", "point -> right -> right -> click -> zoom in"],
            ["Virtual Sorting", "взять предмет, перенести к контейнеру, сбросить и вернуть руку", "point -> click -> right -> click -> left"],
            ["Target Selection", "наведение курсора на цель и подтверждение выбора", "point -> click"],
            ["Object Placement", "якорение объекта на поверхности, масштаб и ориентация", "point -> click -> zoom in -> right"],
            ["Object Inspection", "выбор, поворот, приближение деталей, возврат масштаба", "click -> right -> zoom in -> left -> zoom out"],
            ["Distance Measure", "фиксация двух точек измерения и приближение сегмента", "point -> click -> point -> click -> zoom in"],
            ["Assembly Assist", "выбор детали, переход к слоту, совмещение и фиксация", "point -> click -> right -> point -> click"],
            ["Info Panel", "открытие карточки объекта и перелистывание метаданных", "point -> click -> right -> left -> click"],
            ["Precision Docking", "совмещение объекта с прицелом и подтверждение стыковки", "point -> left -> right -> zoom out -> click"],
            ["Guided Tour", "переход по точкам экскурсии, фокусировка и возврат", "right -> click -> zoom in -> right -> left"],
        ],
        widths=[3.4, 7.1, 5.4],
    )
    add_picture_if_exists(doc, screenshot_sorting, 6.4, "Рисунок 2.4 - Сценарий виртуальной сортировки как пример механики уровня задачи")
    add_para(
        doc,
        "На примере Virtual Sorting видно, как один и тот же словарь жестов используется в более содержательной задаче. Наведение выбирает предмет, нажатие берет его, жест вправо переносит предмет к целевой зоне, повторное нажатие сбрасывает предмет, а жест влево возвращает руку к исходной области. В журнале такой сценарий должен давать не только список меток, но и проверяемую последовательность ожидаемых действий.",
    )

    add_heading(doc, "2.6. Выводы", 2)
    add_para(
        doc,
        "Во второй главе показано, что текущий проект уже вышел за рамки концепта: сформирован словарь жестов, подготовлен эталонный набор на публичных данных, реализованы базовые и временные модели, выделен слой C2 и создан живой интерфейс дополненной реальности. Дополнительная ценность текущего состояния - наличие сценариев, которые превращают распознавание жестов в измеряемые пользовательские задачи.",
    )

    add_heading(doc, "3. ЭКСПЕРИМЕНТАЛЬНАЯ ЧАСТЬ, НОВИЗНА И РЕЗУЛЬТАТЫ", 1)
    add_heading(doc, "3.1. План экспериментов", 2)
    add_para(
        doc,
        "Экспериментальная часть текущего среза построена по принципу постепенного усложнения. Сначала проверяется корректность данных и минимальный проверочный конвейер, затем обучаются модели на малом сбалансированном поднаборе, после этого выполняется полный эталонный прогон на обучающем и тестовом разбиениях IPN, затем проверяется контрольная ветка C1-T с сохранением лучшей эпохи, а после этого система подключается к серверу живого режима, AR-интерфейсу и отчетам по сценариям.",
    )
    add_table(
        doc,
        ["Этап", "Вход", "Выход", "Статус"],
        [
            ["Публичная минимальная проверка", "синтетические ключевые точки", "проверка API и тестового контура", "выполнено"],
            ["Первичный эталон", "175 обучение / 70 тест", "первичное сравнение C0/C1/C1-T", "выполнено"],
            ["Полный эталон", "2405 обучение / 1033 тест", "основные метрики распознавания", "выполнено"],
            ["Проверочная TCN", "полный тест IPN", "контрольная ветка C1-T с общей точностью 0.9071", "выполнено"],
            ["ONNX-экспорт", "C1-T TCN", "ipn_c1t_tcn_full и проверочный ONNX", "выполнено"],
            ["AR-сценарии", "13 сценариев", "конфигурация задач и пользовательский интерфейс", "выполнено"],
            ["Отчет живой задачи", "веб-камера + ONNX", "FPS, задержка, покрытие, предупреждения", "первичный срез выполнен"],
            ["План телефонного AR", "35 локальных роликов", "план целевого домена phone_rear_ar", "ожидает записи видео"],
            ["Мобильный пакет", "метки, предварительная обработка, C2, переносимость", "контракты для мобильного этапа", "выполнено"],
        ],
        widths=[3.2, 4.0, 5.2, 3.0],
    )
    add_para(
        doc,
        "Метрики уровня распознавания включают общую точность, сбалансированную точность, macro F1, weighted F1, матрицу ошибок, медианную и p95 задержку. Для уровня взаимодействия в проекте уже заложены успешность задачи, доля непреднамеренных действий, частота ложных запусков в минуту, точность/полнота действий и число корректировок на задачу; текущие отчеты пока нужно трактовать как инженерную диагностику, потому что пользовательские сценарии с надежной внешней разметкой еще не завершены.",
    )
    add_para(
        doc,
        "В отчете метрики считаются через матрицу ошибок M, где строки соответствуют истинному классу, а столбцы - предсказанному. Такая запись делает сопоставимыми результаты C0, C1, C1-T и проверочной ветки.",
    )
    add_formula_table(
        doc,
        [
            ["(10)", "Accuracy = trace(M) / sum(M)", "Общая доля правильных ответов."],
            ["(11)", "Precision(i)=M_ii/Σ_j M_ji;  Recall(i)=M_ii/Σ_j M_ij", "Точность и полнота отдельного класса."],
            ["(12)", "F1(i)=2·Precision(i)·Recall(i)/(Precision(i)+Recall(i))", "Гармоническое среднее точности и полноты."],
            ["(13)", "MacroF1=(1/K)·Σ_i F1(i);  WeightedF1=Σ_i n_i·F1(i)/Σ_i n_i", "Среднее F1 по классам и с учетом поддержки."],
            ["(14)", "BalancedAccuracy=(1/K)·Σ_i Recall(i)", "Средняя полнота по классам."],
            ["(15)", "L_p95=Q_0.95({l_n});  FPS_mean=N/((t_N-t_1)/1000)", "Задержка p95 и средняя частота кадров."],
        ],
    )

    add_heading(doc, "3.2. Результаты распознавания", 2)
    initial_rows = [
        ["C0 правило", fmt(metrics["C0 initial"]["accuracy"]), fmt(metrics["C0 initial"]["macro_f1"]), fmt(metrics["C0 initial"]["weighted_f1"]), f'{metrics["C0 initial"]["p95_ms"]:.3f}'],
        ["C1 случайный лес", fmt(metrics["C1 RF initial"]["accuracy"]), fmt(metrics["C1 RF initial"]["macro_f1"]), fmt(metrics["C1 RF initial"]["weighted_f1"]), f'{metrics["C1 RF initial"]["p95_ms"]:.3f}'],
        ["C1-T временной прототип", fmt(metrics["C1-T prototype initial"]["accuracy"]), fmt(metrics["C1-T prototype initial"]["macro_f1"]), fmt(metrics["C1-T prototype initial"]["weighted_f1"]), f'{metrics["C1-T prototype initial"]["p95_ms"]:.3f}'],
        ["C1-T компактная TCN", fmt(metrics["C1-T TCN initial"]["accuracy"]), fmt(metrics["C1-T TCN initial"]["macro_f1"]), fmt(metrics["C1-T TCN initial"]["weighted_f1"]), f'{metrics["C1-T TCN initial"]["p95_ms"]:.3f}'],
    ]
    add_table(doc, ["Вариант", "Общая точность", "Macro F1", "Weighted F1", "p95 задержки, мс"], initial_rows, widths=[5.2, 2.4, 2.6, 2.8, 3.0])
    add_para(
        doc,
        "На первичном поднаборе компактная TCN-модель уже показывает лучший результат среди обучаемых вариантов: общая точность 0.7857 и macro F1 0.7705. Случайный лес уступает по точности и имеет существенно большую p95 задержку. Временной прототип остается полезным как легкий проверочный вариант, но не является основным кандидатом для финальной системы.",
    )
    full_rows = [
        ["C0 правило", fmt(metrics["C0 full"]["accuracy"]), fmt(metrics["C0 full"]["balanced_accuracy"]), fmt(metrics["C0 full"]["macro_f1"]), fmt(metrics["C0 full"]["weighted_f1"]), f'{metrics["C0 full"]["p95_ms"]:.3f}'],
        ["C1 случайный лес", fmt(metrics["C1 RF full"]["accuracy"]), fmt(metrics["C1 RF full"]["balanced_accuracy"]), fmt(metrics["C1 RF full"]["macro_f1"]), fmt(metrics["C1 RF full"]["weighted_f1"]), f'{metrics["C1 RF full"]["p95_ms"]:.3f}'],
        ["C1-T компактная TCN", fmt(metrics["C1-T TCN full"]["accuracy"]), fmt(metrics["C1-T TCN full"]["balanced_accuracy"]), fmt(metrics["C1-T TCN full"]["macro_f1"]), fmt(metrics["C1-T TCN full"]["weighted_f1"]), f'{metrics["C1-T TCN full"]["p95_ms"]:.3f}'],
        ["Проверочная C1-T TCN", fmt(metrics["C1-T TCN validated"]["accuracy"]), fmt(metrics["C1-T TCN validated"]["balanced_accuracy"]), fmt(metrics["C1-T TCN validated"]["macro_f1"]), fmt(metrics["C1-T TCN validated"]["weighted_f1"]), f'{metrics["C1-T TCN validated"]["p95_ms"]:.3f}'],
    ]
    add_table(doc, ["Вариант", "Общая точность", "Сбаланс. точность", "Macro F1", "Weighted F1", "p95 задержки, мс"], full_rows, widths=[4.4, 2.1, 2.5, 2.2, 2.4, 2.8])
    add_para(
        doc,
        "На полном эталонном наборе результат меняется качественно. Правило-ориентированный C0 демонстрирует ожидаемо низкую общую точность 0.1820 и macro F1 0.0874. Случайный лес достигает общей точности 0.8955 и weighted F1 0.8930, но его macro F1 остается 0.7987, а p95 задержка составляет 56.409 мс. Компактная C1-T TCN показывает лучший практический баланс: общая точность 0.9061, сбалансированная точность 0.8762, macro F1 0.8504, weighted F1 0.9093 и p95 задержка 3.912 мс. Контрольная проверочная ветка дает немного более высокую общую точность 0.9071 и weighted F1 0.9109, но требует отдельного анализа рисков ложных действий.",
    )
    doc.add_picture(str(chart_path), width=Inches(6.4))
    add_caption(doc, "Рисунок 3.1 - Сравнение вариантов распознавания на полном тестовом разбиении IPN")
    add_picture_if_exists(doc, screenshot_results, 6.4, "Рисунок 3.2 - Экран результатов в локальном интерфейсе Gesture AR")
    per_rows = []
    for label in labels:
        v = metrics["C1-T TCN validated"]["per_class"][label]
        per_rows.append([label, fmt(v["precision"]), fmt(v["recall"]), fmt(v["f1"]), str(int(v["support"]))])
    add_table(doc, ["Класс", "Точность класса", "Полнота", "F1", "Поддержка"], per_rows, widths=[4.0, 3.0, 3.0, 3.0, 2.4])
    add_para(
        doc,
        "По классам в проверочной ветке наиболее сильными являются «указание двумя пальцами» и «нет жеста»: F1 равен 0.9608 и 0.9360 соответственно. По сравнению с предыдущей полной TCN-веткой заметно улучшается жест увеличения масштаба, но ухудшается свайп влево из-за роста числа случаев, когда отсутствие жеста ошибочно превращается в навигационную команду. Это хороший пример того, почему для AR-системы нельзя выбирать модель только по общей точности: направление ложной команды имеет значение для интерфейса.",
    )
    doc.add_picture(str(cm_path), width=Inches(6.3))
    add_caption(doc, "Рисунок 3.3 - Матрица ошибок проверочной C1-T TCN")
    add_para(
        doc,
        "Матрица ошибок показывает, что для проверочной модели ключевой риск сосредоточен не только в близости жестов увеличения и уменьшения масштаба, но и в ложном срабатывании свайпа влево на кадрах без жеста. Поэтому в отчете о рисках дополнительно сравниваются доля ложных действий по классу «нет жеста» и доля ложных свайпов. Текущий анализ показывает, что C1-T full имеет меньший риск ложных свайпов: 7 случаев из 509 фрагментов без жеста, тогда как проверочный вариант дает 28 таких случаев из 509.",
    )

    add_heading(doc, "3.3. Слой взаимодействия и живой диагностический срез", 2)
    add_para(
        doc,
        "Распознавание на сохраненном датасете не гарантирует хорошее AR-взаимодействие. Даже точная модель может вызывать раздражение, если она слишком часто активирует действия, не учитывает контекст сцены или повторяет команду без паузы. Поэтому текущий проект выделяет два режима: прямое управление, где каждый распознанный жест сразу становится действием, и C2 Gate, где команда проходит через фильтр уверенности, устойчивости и времени.",
    )
    add_para(
        doc,
        "Прямое управление полезно для отладки модели и визуальной проверки интерфейса. C2 Gate использует порог уверенности, требование нескольких устойчивых кадров, паузу между повторными действиями и сброс по отсутствию команды. Его задача - сделать поведение системы похожим на намеренное взаимодействие, а не на поток случайных распознаваний.",
    )
    add_table(
        doc,
        ["Параметр C2", "Текущее значение", "Назначение"],
        [
            ["Порог активации", "0.62", "отсекает низкую уверенность"],
            ["Устойчивые кадры", "2", "требует повторного подтверждения жеста"],
            ["Пауза между командами", "250 мс", "подавляет слишком частые повторные действия"],
            ["Сброс по отсутствию жеста", "3 кадра", "возвращает автомат в режим ожидания после отсутствия команды"],
        ],
        widths=[5.2, 3.4, 7.2],
    )
    add_para(
        doc,
        "Показатели взаимодействия в AR-сценариях описываются отдельными формулами, потому что качество пользовательской задачи не сводится к точности распознавания отдельного клипа.",
    )
    add_formula_table(
        doc,
        [
            ["(16)", "FAR_no = FP_action(no_gesture) / N_no_gesture", "Риск ложного действия при отсутствии намеренного жеста; ключевой показатель безопасности интерфейса."],
            ["(17)", "FSR_no = FP_swipe(no_gesture) / N_no_gesture", "Частный риск ложной навигации: в AR такие ошибки особенно заметны для пользователя."],
            ["(18)", "PointerCoverage = N_pointer / N", "Доля кадров, где система смогла построить координату AR-курсора."],
            ["(19)", "RequiredActionCoverage = |A_req ∩ A_obs| / |A_req|", "Покрытие обязательных действий показывает, насколько фактическая сессия соответствует сценарию."],
            ["(20)", "TaskSuccess = 1[RequiredActionCoverage=1 и warnings=∅]", "Бинарное условие успешного выполнения сценария: все обязательные действия есть и нет предупреждений."],
        ],
    )
    add_para(
        doc,
        f"Свежий диагностический срез уровня задачи выполнен в режиме веб-камера + ONNX для сценария размещения объекта. Сессия содержит {task_session.get('frames', 0)} кадров длительностью {task_session.get('duration_ms', 0) / 1000:.2f} с; средняя частота составила {task_session.get('fps', {}).get('mean', 0):.2f} кадра/с, p95 частоты - {task_session.get('fps', {}).get('p95', 0):.2f} кадра/с, среднее время обработки - {task_session.get('processing_ms', {}).get('mean', 0):.3f} мс, p95 обработки - {task_session.get('processing_ms', {}).get('p95', 0):.3f} мс. Средняя доля кадров с обнаруженной рукой равна {task_session.get('detection_rate_mean', 0):.4f}, а покрытие указателя в сценарии составляет {placement_task.get('pointer_coverage', 0):.4f}. Это уже не только проверка запуска камеры, но и рабочий диагностический отчет по задаче.",
    )
    doc.add_picture(str(live_path), width=Inches(6.4))
    add_caption(doc, "Рисунок 3.4 - Живой диагностический срез веб-камера + ONNX")
    add_table(
        doc,
        ["Показатель", "Значение", "Интерпретация"],
        [
            ["Число кадров", str(task_session.get("frames", 0)), "объем кадров в живом журнале"],
            ["Доля обнаружений руки", f'{task_session.get("detection_rate_mean", 0):.4f}', "доля окна, где рука обнаружена"],
            ["Покрытие указателя", f'{placement_task.get("pointer_coverage", 0):.4f}', "доля кадров с координатой AR-курсора"],
            ["Доля бездействия", f'{placement_task.get("idle_ratio", 0):.4f}', "доля кадров без активного действия"],
            ["Покрытие обязательных действий", f'{placement_task.get("required_action_coverage", 0):.4f}', "покрытие обязательных действий сценария"],
            ["Предупреждения", "; ".join(placement_task.get("warnings", []))[:160], "причины, почему сессию нельзя считать успешным пользовательским прогоном"],
        ],
        widths=[4.0, 3.2, 8.4],
    )
    add_para(
        doc,
        "В этой сессии распознавались активные жесты и указатель, однако сценарий не был закрыт успешно: обязательные действия не покрыты полностью. Это не провал системы, а полезная диагностическая точка. Она показывает, что интерфейс уже умеет измерять качество выполнения задачи, а следующий эксперимент должен быть не просто «камера включилась», а строго записанная серия с понятным порядком жестов, внешним контролем разметки и сравнением прямого управления с режимом C2.",
    )
    add_table(
        doc,
        ["Вариант", "False-action rate по no_gesture", "False-swipe rate", "Вывод для AR"],
        [
            ["C1-T full", f'{risk_full["no_gesture_false_action_rate"]:.4f}', f'{risk_full["no_gesture_false_swipe_rate"]:.4f}', "ниже риск ложных навигационных действий"],
            ["Проверочная C1-T", f'{risk_validated["no_gesture_false_action_rate"]:.4f}', f'{risk_validated["no_gesture_false_swipe_rate"]:.4f}', "чуть выше общая точность, но больше ложных свайпов"],
        ],
        widths=[3.5, 4.0, 3.4, 5.0],
    )

    add_heading(doc, "3.4. Перенос на телефонный AR-домен", 2)
    add_para(
        doc,
        "Текущий проект уже подготовлен к переносу на телефонный AR-контур, но этот перенос не следует описывать как завершенный. На Windows отработаны данные, обучение, ONNX-вывод, сервер живого режима и интерфейс. Для телефона зафиксирована другая оболочка выполнения: задняя камера, ARKit/RealityKit для отслеживания мира и отрисовки, тот же принцип извлечения ключевых точек руки, та же предварительная обработка [1,32,74], классификатор Core ML и тот же C2-автомат действий.",
    )
    add_table(
        doc,
        ["Компонент", "Desktop AR", "Phone AR"],
        [
            ["Источник изображения", "веб-камера Windows", "задняя камера телефона"],
            ["Отрисовка AR-сцены", "Three.js", mobile_contract["phone_ar_contract"]["ar_renderer"]],
            ["Вывод модели", "ONNX Runtime / PyTorch", "Core ML, конвертация отложена"],
            ["Вход модели", "окно [1,32,74]", str(preprocess_contract["model_input_shape"])],
            ["Контекстный слой", "C2", "то же соответствие «метка -> действие» и те же пороги"],
            ["Назначение локального домена", "инженерная проверка веб-камеры", "адаптация к телефонному AR-домену"],
        ],
        widths=[4.0, 5.4, 6.2],
    )
    add_para(
        doc,
        f"Отчет готовности доменов фиксирует {domain_report['by_capture_domain'].get('ipn_hand', 0)} записей в домене IPN Hand и {domain_report['by_capture_domain'].get('phone_rear_ar', 0)} планируемых записей в домене phone_rear_ar. Технический статус переноса: {domain_report.get('domain_transfer_status', '')}. Все {domain_report['local_phone']['planned_records']} локальных записей пока ожидают появления исходных видео, поэтому текущий мобильный этап корректно называть подготовленным контрактом, а не завершенной валидацией.",
    )
    add_para(
        doc,
        "Такое разделение важно для научной честности. Если сразу смешать IPN Hand и локальные ролики, станет трудно понять, что именно улучшило результат: общая устойчивость признаков, дообучение на конкретной руке, подбор порога C2 или случайное совпадение условий съемки. Поэтому план переноса задает строгую последовательность: публичная модель на IPN, проверка без дообучения на phone_rear_ar, калибровка C2, затем при необходимости тонкая настройка TCN.",
    )
    add_para(
        doc,
        "Для будущего телефонного этапа удобно заранее задать формулы переносимости. Они позволяют отделить собственно качество модели от различий камеры, освещения, масштаба руки и способа отрисовки AR-сцены.",
    )
    add_formula_table(
        doc,
        [
            ["(21)", "Z_phone = φ(X_phone),  Z_phone ∈ R^(1×32×74)", "Та же функция предварительной обработки φ должна давать совместимый вход для Core ML-модели."],
            ["(22)", "Gap_D(m) = |m_IPN - m_phone|", "Доменный разрыв по метрике m показывает, насколько падает качество при переходе к телефонной камере."],
            ["(23)", "τ* = argmin_τ [α·FAR(τ) + β·MissRate(τ)]", "Калибровка C2 может быть оформлена как поиск порога, балансирующего ложные действия и пропуски команд."],
        ],
    )

    add_heading(doc, "3.5. Научная новизна и ограничения", 2)
    add_para(
        doc,
        "На текущем этапе новизна работы формулируется как связка временного распознавания по ключевым точкам кисти, контекстно-зависимой политики действий, первичного публичного эталона и переносимого AR-контура. В отличие от работ, где распознавание жестов рассматривается изолированно, здесь классификатор сразу помещается в слой действий и оценивается с учетом отсутствия жеста, устойчивости, паузы между командами и метрик уровня пользовательской задачи.",
    )
    add_para(
        doc,
        "Первая составляющая новизны - единый контракт данных для публичных и локальных записей. Все клипы приводятся к одинаковому формату [T,21,3] + маска + уверенность, что позволяет сравнивать модели и выполнять локальную адаптацию без смены архитектуры.",
    )
    add_para(
        doc,
        "Вторая составляющая - разделение распознавания и взаимодействия. Accuracy и macro F1 фиксируют качество классификатора, но пользовательский результат зависит от того, какие действия реально дошли до AR-сцены. Поэтому C2 рассматривается как самостоятельный исследуемый компонент, а не как техническая мелочь интерфейса.",
    )
    add_para(
        doc,
        "Третья составляющая - воспроизводимость. Проект уже содержит CLI, тесты, артефакты, отчеты, ONNX-экспорт, мобильный bundle и сценарии интерфейса. Это снижает риск, что результат диссертации окажется только демонстрацией на одной машине без возможности повторения экспериментов.",
    )
    add_para(
        doc,
        "Ограничения также нужно фиксировать честно. Контрольная проверочная ветка уже появилась, но она не снимает вопрос выбора модели: по общей точности она чуть лучше, а по риску ложных свайпов хуже. C2 пока не прошел полноценную пользовательскую оценку, локальные телефонные клипы еще не записаны, переносимость Core ML не доведена до финальной iOS-сборки, а отчет живой задачи пока показывает диагностическую готовность, но не финальное качество пользовательского взаимодействия.",
    )

    add_heading(doc, "ЗАКЛЮЧЕНИЕ", 1)
    add_para(
        doc,
        "В ходе текущего этапа выполнен переход от концептуального описания бесконтактного AR-взаимодействия к работающему исследовательскому конвейеру. Подготовлен публичный эталонный набор на основе IPN Hand, реализовано единое представление ключевых точек кисти, обучены и сравнены варианты C0, C1 и C1-T, создан контекстный слой C2, выполнен ONNX-экспорт, запущен live AR-контур и добавлены сценарии пользовательских задач.",
    )
    add_para(
        doc,
        "Главный экспериментальный результат текущего среза - компактная C1-T TCN на полном тестовом разбиении IPN достигает общей точности 0.9061, macro F1 0.8504 и p95 задержки 3.912 мс, а проверочный вариант достигает общей точности 0.9071 и weighted F1 0.9109. При этом отчет о рисках показывает, что выбор модели должен учитывать не только общую точность, но и структуру ложных действий, особенно ложные свайпы в классе «нет жеста».",
    )
    add_para(
        doc,
        "Научная рамка магистерской работы уточнена: исследование должно доказывать не только повышение точности распознавания, но и влияние временной модели и контекстной политики действий на практические метрики выполнения задач. Поэтому следующими шагами являются запись 35 локальных роликов phone_rear_ar, проверка публичной модели на телефонном домене без дообучения, калибровка C2, воспроизведение сценариев по заданному скрипту и пользовательская оценка прямого управления против режима C2 Gate.",
    )
    add_para(
        doc,
        "Документ не является финальной диссертацией, но служит рабочим фиксатором состояния проекта и может использоваться как основа для последующей полной магистерской редакции.",
    )

    add_literature(doc)

    # Ensure every section after the title page keeps the same margins and page numbering.
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    doc.core_properties.author = "Южаков М.В."
    doc.core_properties.title = "Рабочий отчет магистерской диссертации по проекту распознавания жестов"
    doc.core_properties.subject = "Контекстно-зависимое распознавание жестов по ключевым точкам кисти для дополненной реальности"
    doc.core_properties.keywords = "дополненная реальность, жесты, MediaPipe, IPN Hand, TCN, ONNX"
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
    print(DOCX_PATH)
