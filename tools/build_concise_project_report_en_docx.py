from __future__ import annotations

import csv
import json
import math
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "generated"
FIG_DIR = OUT_DIR / "concise_project_report_assets"
DOCX_PATH = OUT_DIR / "Yuzhakov_MV_Concise_Project_Report_2026-05-26_EN.docx"

ACCENT = RGBColor(46, 116, 181)
ACCENT_DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 96, 106)
HEADER_FILL = "F2F4F7"
BLUE_FILL = "E8EEF5"
GREEN_FILL = "E2F0D9"
GOLD_FILL = "FFF2CC"
GRAY_FILL = "F7F9FC"
BORDER = "C8CED8"


def read_json(relative: str) -> dict:
    return json.loads((ROOT / relative).read_text(encoding="utf-8"))


def read_csv(relative: str) -> list[dict[str, str]]:
    with (ROOT / relative).open(encoding="utf-8", newline="") as fh:
        return list(csv.DictReader(fh))


def fmt(value: str | float | int | None, digits: int = 4) -> str:
    if value is None or value == "":
        return "-"
    if isinstance(value, str):
        try:
            value = float(value)
        except ValueError:
            return value
    return f"{float(value):.{digits}f}"


def fmt_pct(value: str | float | int | None, digits: int = 1) -> str:
    if value is None or value == "":
        return "-"
    return f"{float(value) * 100:.{digits}f}%"


def set_run_font(run, name: str = "Calibri", size: int | None = None, color: RGBColor | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        element = tc_mar.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: int = 9,
    center: bool = False,
    fill: str | None = None,
    color: RGBColor | None = None,
) -> None:
    cell.text = ""
    set_cell_margins(cell)
    if fill:
        set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = para.add_run(text)
    run.bold = bold
    set_run_font(run, size=size, color=color)


def set_table_borders(table, color: str = BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths_cm: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)
    tbl_grid = table._tbl.tblGrid
    if tbl_grid is not None:
        for child in list(tbl_grid):
            tbl_grid.remove(child)
        for width in widths_cm:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(int(width / 2.54 * 1440)))
            tbl_grid.append(col)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        hdr = OxmlElement("w:tblHeader")
        hdr.set(qn("w:val"), "true")
        tr_pr.append(hdr)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_cm: list[float],
    *,
    header_fill: str = HEADER_FILL,
    font_size: int = 9,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    set_repeat_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=font_size, center=True, fill=header_fill)
    for values in rows:
        row = table.add_row()
        set_row_cant_split(row)
        for idx, value in enumerate(values):
            center = idx > 0 and len(str(value)) < 18
            set_cell_text(row.cells[idx], str(value), size=font_size, center=center)
    set_table_width(table, widths_cm)
    doc.add_paragraph()


def add_key_value_table(doc: Document, rows: list[tuple[str, str]], *, fill: str = GRAY_FILL) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for label, value in rows:
        row = table.add_row()
        set_row_cant_split(row)
        set_cell_text(row.cells[0], label, bold=True, size=9, fill=fill)
        set_cell_text(row.cells[1], value, size=9)
    set_table_width(table, [4.1, 12.0])
    doc.add_paragraph()


def add_callout(doc: Document, title: str, body: str, *, fill: str = BLUE_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, "AEBAD0")
    set_row_cant_split(table.rows[0])
    cell = table.rows[0].cells[0]
    cell.text = ""
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    set_cell_shading(cell, fill)
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(title)
    run.bold = True
    set_run_font(run, size=10, color=ACCENT_DARK)
    para2 = cell.add_paragraph()
    para2.paragraph_format.space_after = Pt(0)
    run2 = para2.add_run(body)
    set_run_font(run2, size=10)
    set_table_width(table, [16.1])
    doc.add_paragraph()


def add_placeholder(doc: Document, number: int, instruction: str) -> None:
    title = f"Screenshot placeholder {number}"
    body = instruction + " Replace this box with your own image after running the webcam demo."
    add_callout(doc, title, body, fill=GOLD_FILL)


def add_caption(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(8)
    run = para.add_run(text)
    run.italic = True
    set_run_font(run, size=9, color=MUTED)


def add_figure(doc: Document, path: Path, caption: str, *, width_in: float = 6.1) -> None:
    if not path.exists():
        add_callout(doc, "Missing figure", f"Expected figure was not found: {path}", fill=GOLD_FILL)
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    add_caption(doc, caption)


def get_font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def wrap(draw: ImageDraw.ImageDraw, text: str, font, width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = word if not current else current + " " + word
        box = draw.textbbox((0, 0), candidate, font=font)
        if box[2] - box[0] <= width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_box(draw: ImageDraw.ImageDraw, rect: tuple[int, int, int, int], text: str, fill: str, outline: str) -> None:
    draw.rounded_rectangle(rect, radius=16, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = rect
    font = get_font(24, True)
    small = get_font(19)
    lines = wrap(draw, text, font, x2 - x1 - 36)
    if len(lines) > 2:
        font = small
        lines = wrap(draw, text, font, x2 - x1 - 36)
    line_h = font.size + 7 if hasattr(font, "size") else 26
    start_y = y1 + (y2 - y1 - line_h * len(lines)) / 2
    for line in lines:
        box = draw.textbbox((0, 0), line, font=font)
        draw.text((x1 + (x2 - x1 - (box[2] - box[0])) / 2, start_y), line, fill="#1F2937", font=font)
        start_y += line_h


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#2E74B5") -> None:
    draw.line([start, end], fill=color, width=5)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 16
    points = [
        end,
        (end[0] - size * math.cos(angle - math.pi / 6), end[1] - size * math.sin(angle - math.pi / 6)),
        (end[0] - size * math.cos(angle + math.pi / 6), end[1] - size * math.sin(angle + math.pi / 6)),
    ]
    draw.polygon(points, fill=color)


def create_architecture_diagram() -> Path:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    path = FIG_DIR / "c4_webcam_architecture.png"
    image = Image.new("RGB", (1800, 820), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    title_font = get_font(38, True)
    note_font = get_font(22)
    draw.text((60, 36), "Current system pipeline: webcam AR interaction with action-safe control", fill="#0B2545", font=title_font)
    draw.text(
        (60, 90),
        "The current research value is the full chain from hand landmarks to task-aware AR actions, not only isolated gesture classification.",
        fill="#4B5563",
        font=note_font,
    )
    boxes = [
        ((80, 190, 390, 330), "Webcam / planned phone rear camera", "#E8EEF5"),
        ((490, 190, 800, 330), "Hand landmarks and temporal window", "#F7F9FC"),
        ((900, 190, 1210, 330), "C1-T temporal recognizer", "#E2F0D9"),
        ((1310, 190, 1640, 330), "C4 task-aware safety gate", "#FFF2CC"),
        ((900, 510, 1210, 650), "Backend WebSocket session", "#F7F9FC"),
        ((1310, 510, 1640, 650), "React / Three.js AR task interface", "#E8EEF5"),
    ]
    for rect, text, fill in boxes:
        draw_box(draw, rect, text, fill, "#8CA7C8")
    draw_arrow(draw, (390, 260), (490, 260))
    draw_arrow(draw, (800, 260), (900, 260))
    draw_arrow(draw, (1210, 260), (1310, 260))
    draw_arrow(draw, (1475, 330), (1475, 510))
    draw_arrow(draw, (1310, 580), (1210, 580))
    draw_arrow(draw, (1055, 510), (1055, 330))
    draw.text((105, 370), "Input: RGB frames -> 21 hand landmarks", fill="#4B5563", font=note_font)
    draw.text((925, 370), "Output: gesture label + confidence", fill="#4B5563", font=note_font)
    draw.text((1328, 370), "Output: allowed AR action or idle", fill="#4B5563", font=note_font)
    draw.text((925, 690), "Live metrics: FPS, processing time, detection rate, action stream", fill="#4B5563", font=note_font)
    image.save(path)
    return path


def create_ui_snapshot_crop() -> Path:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    source = ROOT / "artifacts" / "screenshots" / "ar_interface_virtual_sorting_final_checked.png"
    target = FIG_DIR / "ar_interface_virtual_sorting_context_crop.png"
    if not source.exists():
        return source
    with Image.open(source).convert("RGB") as image:
        width, height = image.size
        crop_bottom = min(height, 640)
        cropped = image.crop((0, 0, width, crop_bottom))
        cropped.save(target)
    return target


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 16, 8),
        ("Heading 2", 13, ACCENT, 12, 6),
        ("Heading 3", 12, ACCENT_DARK, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run("Current project snapshot - generated 26 May 2026")
    set_run_font(run, size=9, color=MUTED)


def add_title(doc: Document) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run("Gesture AR Project: Concise Current Report")
    run.bold = True
    set_run_font(run, size=24, color=RGBColor(11, 37, 69))

    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(12)
    run = para.add_run("C4 task-aware, webcam-capable AR interaction snapshot")
    set_run_font(run, size=13, color=MUTED)

    add_key_value_table(
        doc,
        [
            ("Project", "Hand-gesture control for an augmented-reality web interface."),
            ("Snapshot date", date(2026, 5, 26).strftime("%d %B %Y")),
            ("Current emphasis", "Gesture recognition plus action-safe control inside full AR task scenarios."),
            ("Demo status", "The system can be launched locally and interacted with through a webcam; screenshot slots are marked below."),
        ],
    )


def add_section_1(doc: Document, arch_path: Path) -> None:
    doc.add_heading("1. Project Essence", level=1)
    p = doc.add_paragraph()
    p.add_run("The project develops a gesture-driven augmented-reality interaction system. ").bold = True
    p.add_run(
        "A live camera stream is processed into hand landmarks, a temporal model recognizes command gestures, "
        "and a web AR interface converts those commands into actions such as pointing, confirming, navigation and zooming. "
        "The practical goal is a hands-free control loop for AR tasks, where the user can manipulate virtual objects "
        "without touching the keyboard, mouse or screen."
    )
    doc.add_paragraph(
        "The current stage is no longer limited to measuring a classifier in isolation. The new C4 layer evaluates "
        "whether recognized gestures should become interface actions in the current task context. This is important "
        "because false actions in AR can be more harmful than a wrong label: an accidental confirmation, navigation step "
        "or zoom may move the scenario into an incorrect state."
    )
    add_callout(
        doc,
        "Main contribution in the current snapshot",
        "A temporal gesture recognizer is combined with a task-aware safety gate. The gate keeps task completion near "
        "the previous C3+C2 baseline while reducing unintended and costly actions in full AR scenarios.",
        fill=BLUE_FILL,
    )
    add_figure(doc, arch_path, "Figure 1. System-level pipeline used in the current C4 snapshot.", width_in=6.25)

    doc.add_heading("Core terminology and formulas", level=2)
    doc.add_paragraph(
        "The formulas below are intentionally compact; they define the main quantities used by the implementation and "
        "the current evaluation reports."
    )
    add_table(
        doc,
        ["Formula", "Meaning"],
        [
            ["X_t in R^(T x 21 x 3)", "Temporal window of 21 hand landmarks with three coordinates per landmark."],
            ["Z_t = phi(X_t)", "Feature tensor after normalization, motion features and temporal formatting."],
            ["p_t = softmax(f_theta(Z_t)); y_t = argmax p_t", "TCN gesture prediction and confidence distribution."],
            ["a_t = g(y_t, p_t, s_t, c_t)", "Action decision using label, confidence, stability state and task context."],
            ["UAR = FP_action / (TP_action + FP_action)", "Unintended action rate used to measure interaction safety."],
            ["FC = sum_i cost(a_i) * I(false_i) / sum_i cost(a_i)", "False action cost rate; risky actions receive higher penalty."],
        ],
        [6.1, 9.9],
        header_fill=GREEN_FILL,
        font_size=9,
    )
    formula_table = doc.tables[-1]
    for row in formula_table.rows[1:]:
        for run in row.cells[0].paragraphs[0].runs:
            set_run_font(run, name="Cambria Math", size=9)


def add_section_2(doc: Document, data: dict) -> None:
    doc.add_heading("2. Current State of the System", level=1)
    doc.add_paragraph(
        "The current repository contains a working research pipeline, a FastAPI backend, a React/Three.js AR frontend, "
        "offline recognition reports, C4 task-level benchmarks and a one-command local demo script. The frontend can "
        "switch to Camera Stream mode and receive live actions from the backend through WebSocket."
    )
    doc.add_paragraph(
        "Typical launch command: powershell -ExecutionPolicy Bypass -File .\\scripts\\start_ar_demo.ps1. "
        "After the interface opens, the user selects an AR task, presses Start Task, allows webcam access and performs "
        "the highlighted gesture sequence."
    )

    rec = data["recognition"]["recognition"]
    lat = data["recognition"]["latency"]
    live = data["live"]["session"]
    domain = data["domain"]
    c4 = data["c4_summary"]
    direct = c4["c1t_direct"]
    c3_default = c4["c3_c2_default"]
    task_aware = c4["c4_task_aware"]

    add_table(
        doc,
        ["Area", "Current measurement", "Interpretation"],
        [
            ["Recognizer", f"Accuracy {fmt_pct(rec['accuracy'])}; macro F1 {fmt(rec['macro_f1'])}", "Validated C1-T temporal model is suitable as the live recognizer baseline."],
            ["Offline latency", f"Median {fmt(lat['offline_latency_ms_median'], 2)} ms; p95 {fmt(lat['offline_latency_ms_p95'], 2)} ms", "Model inference is fast enough for interactive use on pre-extracted landmarks."],
            ["Webcam demo", f"Mean FPS {fmt(live['fps']['mean'], 2)}; processing p95 {fmt(live['processing_ms']['p95'], 2)} ms", "Existing live report confirms webcam execution, although the task was not fully completed in that sample."],
            ["C4 task-aware", f"Precision {fmt(task_aware['action_precision_mean'])}; UAR {fmt(task_aware['unintended_action_rate_mean'])}", "Task context reduces accidental actions while keeping the task-success level of C3+C2 default."],
            ["False action cost", f"C1-T direct {fmt(direct['false_action_cost_rate_mean'])}; C3+C2 {fmt(c3_default['false_action_cost_rate_mean'])}; C4 task-aware {fmt(task_aware['false_action_cost_rate_mean'])}", "The task-aware gate cuts costly false actions compared with direct recognition and the C3+C2 policy."],
            ["Domain transfer", f"{domain['local_phone']['planned_records']} phone-rear clips planned; {domain['local_phone']['missing_raw_video_count']} still missing", "Phone rear-camera validation is planned but not yet recorded."],
        ],
        [3.8, 5.3, 6.9],
        header_fill=HEADER_FILL,
        font_size=8,
    )

    add_figure(
        doc,
        ROOT / "artifacts" / "figures" / "c4_task_false_action_cost_rate.png",
        "Figure 2. Task-level false action cost rate across interaction policies.",
        width_in=5.8,
    )
    add_figure(
        doc,
        create_ui_snapshot_crop(),
        "Figure 3. Existing AR interface screenshot from the current demo assets.",
        width_in=4.8,
    )

    doc.add_heading("Webcam screenshot slots", level=2)
    doc.add_paragraph(
        "The report already documents that the system supports webcam interaction. For the next hand-prepared version, "
        "insert fresh screenshots from your own run at the marked places below."
    )
    add_placeholder(
        doc,
        1,
        "Insert a screenshot of the AR task running with Camera Stream visible, a virtual object or scenario overlay, "
        "and the Live Status block showing FPS, detection and current action.",
    )
    add_placeholder(
        doc,
        2,
        "Insert a screenshot of Advanced Controls where the selected interaction mode is visible, preferably C4 task-aware "
        "or the mode used for the experiment.",
    )
    add_placeholder(
        doc,
        3,
        "Optional: insert a screenshot after a short completed or attempted task, showing the task panel, current step "
        "state and any action counters or status messages.",
    )


def add_section_3(doc: Document) -> None:
    doc.add_heading("3. Further Development and Expected Final Result", level=1)
    doc.add_paragraph(
        "The final master project should present the system as a complete AR interaction stack: dataset preparation, "
        "temporal gesture recognition, action-safe control, live webcam demonstration and a measured evaluation across "
        "both offline and task-level conditions."
    )
    add_table(
        doc,
        ["Work item", "Purpose", "Expected final evidence"],
        [
            ["Record local phone-rear clips", "Check whether the model survives the target AR capture geometry.", "Filled local_phone manifest, extracted landmarks and domain-transfer metrics."],
            ["Improve task-aware C4 policy", "Balance safety and responsiveness for each AR scenario.", "Ablation table comparing direct, C3+C2, C4 safety and C4 task-aware modes."],
            ["Run repeated webcam sessions", "Show that the interface works beyond a single sample run.", "Several screenshots plus FPS, latency, detection and completion statistics."],
            ["Prepare mobile/export path", "Move from desktop prototype toward deployable interaction.", "Updated ONNX/Core ML bundle and notes on mobile constraints."],
            ["Finalize thesis narrative", "Connect implementation, experiments, novelty and limitations.", "Three-chapter thesis draft with figures, formulas and final result tables."],
        ],
        [4.2, 5.8, 6.0],
        header_fill=BLUE_FILL,
        font_size=8,
    )
    add_callout(
        doc,
        "Expected final form",
        "A webcam-capable and phone-validated AR gesture-control prototype with a documented recognition model, "
        "a risk-aware action policy, repeatable scenario benchmarks and visual proof of live interaction.",
        fill=GREEN_FILL,
    )

    doc.add_heading("Short conclusion", level=2)
    doc.add_paragraph(
        "At this project state, the essential idea is already demonstrable: a user can interact with the AR interface "
        "through a camera stream, while the research pipeline measures not only gesture labels but also the safety of "
        "resulting interface actions. The strongest current result is the C4 task-aware policy: it reduces unintended "
        "and costly actions while preserving approximately the same task-success level as the previous guarded baseline. "
        "The main remaining work is to collect the planned phone-rear data, repeat webcam demonstrations and tune the "
        "final balance between responsiveness and safety."
    )


def build_report() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    arch = create_architecture_diagram()
    c4_rows = read_csv("artifacts/reports/c4_task_tables/c4_task_summary.csv")
    c4_summary = {row["method"]: row for row in c4_rows}
    data = {
        "recognition": read_json("artifacts/reports/ipn_c1t_tcn_full_validated_recognition.json"),
        "live": read_json("artifacts/reports/live_task_report.json"),
        "domain": read_json("artifacts/reports/domain_readiness.json"),
        "c4_summary": c4_summary,
    }

    doc = Document()
    style_document(doc)
    add_title(doc)
    add_section_1(doc, arch)
    add_section_2(doc, data)
    add_section_3(doc)
    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build_report())
