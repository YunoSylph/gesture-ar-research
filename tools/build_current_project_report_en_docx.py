from __future__ import annotations

import csv
import json
import math
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "generated"
ASSET_DIR = OUT_DIR / "current_project_report_assets"
DOCX_PATH = OUT_DIR / "Yuzhakov_MV_Current_Project_Report_2026-05-28_EN.docx"

ACCENT = RGBColor(31, 77, 120)
ACCENT_DARK = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(90, 96, 106)
BORDER = "C8CED8"
HEADER_FILL = "F2F4F7"
BLUE_FILL = "E8EEF5"
GREEN_FILL = "EEF5EA"
GOLD_FILL = "FFF2CC"
GRAY_FILL = "F7F9FC"


def read_json(relative: str) -> dict:
    return json.loads((ROOT / relative).read_text(encoding="utf-8"))


def read_csv(relative: str) -> list[dict[str, str]]:
    with (ROOT / relative).open(encoding="utf-8", newline="") as handle:
        return list(csv.DictReader(handle))


def fmt(value: str | float | int | None, digits: int = 3) -> str:
    if value is None or value == "":
        return "-"
    if isinstance(value, str):
        try:
            value = float(value)
        except ValueError:
            return value
    return f"{float(value):.{digits}f}"


def fmt_pct(value: str | float | int | None, digits: int = 1) -> str:
    if value is None or value == "":
        return "-"
    return f"{float(value) * 100:.{digits}f}%"


def set_run_font(run, name: str = "Calibri", size: int | None = None, color: RGBColor | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        element = tc_mar.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color: str = BORDER) -> None:
    borders = table._tbl.tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table._tbl.tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: int = 9,
    center: bool = False,
    fill: str | None = None,
    color: RGBColor | None = None,
) -> None:
    cell.text = ""
    set_cell_margins(cell)
    if fill:
        set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = paragraph.add_run(text)
    run.bold = bold
    set_run_font(run, size=size, color=color)


def set_table_width(table, widths_cm: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths_cm):
            row.cells[index].width = Cm(width)
    grid = table._tbl.tblGrid
    if grid is not None:
        for child in list(grid):
            grid.remove(child)
        for width in widths_cm:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(int(width / 2.54 * 1440)))
            grid.append(col)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_cm: list[float],
    *,
    header_fill: str = HEADER_FILL,
    font_size: int = 9,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    set_repeat_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, center=True, fill=header_fill, size=font_size)
    for values in rows:
        row = table.add_row()
        set_row_cant_split(row)
        for index, value in enumerate(values):
            set_cell_text(row.cells[index], str(value), size=font_size, center=index > 0 and len(str(value)) < 20)
    set_table_width(table, widths_cm)
    doc.add_paragraph()


def add_key_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for label, value in rows:
        row = table.add_row()
        set_row_cant_split(row)
        set_cell_text(row.cells[0], label, bold=True, fill=GRAY_FILL, size=9)
        set_cell_text(row.cells[1], value, size=9)
    set_table_width(table, [4.2, 11.8])
    doc.add_paragraph()


def add_callout(doc: Document, title: str, body: str, *, fill: str = BLUE_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, "AEBAD0")
    set_row_cant_split(table.rows[0])
    cell = table.rows[0].cells[0]
    cell.text = ""
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    set_cell_shading(cell, fill)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(title)
    r1.bold = True
    set_run_font(r1, size=10, color=ACCENT_DARK)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10)
    set_table_width(table, [16.1])
    doc.add_paragraph()


def add_placeholder(doc: Document, number: int, instruction: str) -> None:
    add_callout(doc, f"Live screenshot placeholder {number}", instruction, fill=GOLD_FILL)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.italic = True
    set_run_font(run, size=9, color=MUTED)


def add_figure(doc: Document, path: Path, caption: str, *, width_in: float = 6.1) -> None:
    if not path.exists():
        add_callout(doc, "Missing figure", f"Expected image was not found: {path}", fill=GOLD_FILL)
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(path), width=Inches(width_in))
    add_caption(doc, caption)


def get_font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def create_bar_chart(
    filename: str,
    title: str,
    values: list[tuple[str, float]],
    *,
    color: str = "#2E74B5",
    max_value: float | None = None,
    lower_is_better: bool = False,
) -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    path = ASSET_DIR / filename
    width, height = 1500, 760
    image = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    title_font = get_font(38, True)
    label_font = get_font(22)
    small_font = get_font(18)
    draw.text((70, 48), title, fill="#0B2545", font=title_font)
    chart_left, chart_top = 130, 145
    chart_right, chart_bottom = width - 70, height - 150
    draw.line((chart_left, chart_bottom, chart_right, chart_bottom), fill="#9AA5B1", width=2)
    draw.line((chart_left, chart_top, chart_left, chart_bottom), fill="#9AA5B1", width=2)
    top = max_value or max(v for _, v in values) * 1.12
    bar_area = chart_right - chart_left
    gap = 34
    bar_width = (bar_area - gap * (len(values) + 1)) / len(values)
    for i, (label, value) in enumerate(values):
        x1 = chart_left + gap + i * (bar_width + gap)
        x2 = x1 + bar_width
        y1 = chart_bottom - (value / top) * (chart_bottom - chart_top)
        draw.rounded_rectangle((x1, y1, x2, chart_bottom), radius=8, fill=color)
        draw.text((x1 + 4, y1 - 30), f"{value:.3f}", fill="#0B2545", font=small_font)
        words = label.split()
        label_text = "\n".join(words) if len(label) > 13 else label
        draw.multiline_text((x1 + 4, chart_bottom + 18), label_text, fill="#111827", font=small_font, spacing=3)
    direction = "lower is better" if lower_is_better else "higher is better"
    draw.text((70, height - 58), direction, fill="#4B5563", font=label_font)
    image.save(path)
    return path


def create_tradeoff_chart(summary: dict[str, dict[str, str]]) -> Path:
    rows = [
        ("M1 Baseline", float(summary["baseline_direct"]["false_action_cost_rate_mean"]), float(summary["baseline_direct"]["task_success_rate_mean"])),
        ("M2 Robust C6", float(summary["robust_recognizer_direct"]["false_action_cost_rate_mean"]), float(summary["robust_recognizer_direct"]["task_success_rate_mean"])),
        ("M3 TARC", float(summary["proposed_tarc"]["false_action_cost_rate_mean"]), float(summary["proposed_tarc"]["task_success_rate_mean"])),
    ]
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    path = ASSET_DIR / "official_method_tradeoff.png"
    image = Image.new("RGB", (1500, 760), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    title_font = get_font(38, True)
    label_font = get_font(22)
    small_font = get_font(18)
    draw.text((70, 48), "Official replay benchmark: safety and completion", fill="#0B2545", font=title_font)
    left, top, right, bottom = 130, 145, 1430, 610
    draw.line((left, bottom, right, bottom), fill="#9AA5B1", width=2)
    draw.line((left, top, left, bottom), fill="#9AA5B1", width=2)
    group_width = (right - left) / len(rows)
    max_false = max(r[1] for r in rows) * 1.25
    max_success = 0.65
    for i, (label, false_cost, success) in enumerate(rows):
        gx = left + i * group_width + 55
        false_h = (false_cost / max_false) * (bottom - top)
        succ_h = (success / max_success) * (bottom - top)
        draw.rounded_rectangle((gx, bottom - false_h, gx + 130, bottom), radius=8, fill="#D97706")
        draw.rounded_rectangle((gx + 170, bottom - succ_h, gx + 300, bottom), radius=8, fill="#2E74B5")
        draw.text((gx, bottom - false_h - 28), f"{false_cost:.3f}", fill="#7A3E00", font=small_font)
        draw.text((gx + 170, bottom - succ_h - 28), f"{success:.3f}", fill="#0B2545", font=small_font)
        draw.text((gx, bottom + 18), label, fill="#111827", font=small_font)
    draw.rounded_rectangle((970, 80, 1430, 125), radius=8, fill="#FFF2CC", outline="#D97706")
    draw.text((985, 90), "orange = false action cost; blue = task success", fill="#111827", font=small_font)
    draw.text((70, 690), "TARC lowers false action cost while keeping task success at the baseline level.", fill="#4B5563", font=label_font)
    image.save(path)
    return path


def crop_screenshot(source: Path, target_name: str) -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    target = ASSET_DIR / target_name
    if not source.exists():
        return source
    with Image.open(source).convert("RGB") as image:
        # Keep the exact wide viewport. Resize only to reduce DOCX weight.
        image.thumbnail((1500, 850), Image.Resampling.LANCZOS)
        image.save(target, quality=94)
    return target


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 16, 8),
        ("Heading 2", 13, ACCENT, 12, 6),
        ("Heading 3", 12, ACCENT_DARK, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run("Current academic snapshot - generated 28 May 2026")
    set_run_font(run, size=9, color=MUTED)


def add_title(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run("Gesture AR Project: Current Academic Snapshot")
    run.bold = True
    set_run_font(run, size=23, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("C6 recognition plus task-aware risk-calibrated AR control")
    set_run_font(run, size=13, color=MUTED)

    add_key_value_table(
        doc,
        [
            ("Snapshot date", "28 May 2026"),
            ("Current prototype", "FastAPI live backend, React/Three.js AR interface, webcam mode, C6 Ensemble recognizer and TARC controller."),
            ("Research claim", "A guided AR interaction pipeline can reduce costly unintended commands without giving up baseline-level task completion."),
            ("Evidence status", "Offline public-data recognition and replay-based AR task benchmark are ready; live webcam screenshots and repeated user sessions still need to be inserted."),
        ],
    )


def add_section_work_description(doc: Document) -> None:
    doc.add_heading("1. Work Description", level=1)
    doc.add_paragraph(
        "The project is a gesture-controlled augmented-reality interaction system. It converts a webcam stream into hand "
        "landmarks, recognizes command gestures, and applies a task-aware controller before the AR interface executes an action."
    )
    doc.add_paragraph(
        "The current version is best described as a full interaction pipeline rather than a standalone gesture classifier. "
        "The system includes dataset manifests, temporal recognition models, calibrated ensemble fusion, risk-cost action "
        "filtering, scripted AR tasks, a local live backend and a web-based AR-style interface."
    )
    add_callout(
        doc,
        "Academic novelty, stated conservatively",
        "The original contribution is not a new universal neural architecture. It is the integration of a robust gesture "
        "recognizer with a task-aware risk controller for guided AR workflows, evaluated with action cost and task completion "
        "rather than only classification accuracy.",
        fill=BLUE_FILL,
    )
    add_callout(
        doc,
        "Boundary of the claim",
        "The current evidence supports public-dataset recognition and replay-based interaction safety. It does not yet prove "
        "general open-world AR control or phone rear-camera transfer; those require live session logs and local video data.",
        fill=GOLD_FILL,
    )
    doc.add_page_break()
    doc.add_heading("Methodological model and formulas", level=2)
    doc.add_paragraph(
        "The method treats recognition as one component of a control loop. A gesture prediction becomes only an action "
        "proposal; the TARC controller may accept it, delay it, or replace it with idle."
    )
    add_table(
        doc,
        ["Formula", "Role in the project"],
        [
            ["X_t in R^(T x 21 x 3), T = 32", "Temporal landmark window from one hand: 21 landmarks, three coordinates each."],
            ["z_k = f_k(X_t)", "Logits from recognizer k. C6 uses the validated TCN and an augmented TCN."],
            ["p_C6 = softmax((mean_k z_k + alpha h_geom + b_y) / tau)", "Calibrated ensemble fusion with a lightweight geometry-aware correction."],
            ["y_t = argmax_y p_C6(y | X_t)", "Current gesture label used as an action proposal, not as a direct command."],
            ["a_t = pi(y_t) if p_t >= theta_eff and stable(y_t) and cooldown_ok else idle", "TARC action gate based on confidence, temporal stability and cooldown."],
            ["theta_eff(a,c) = theta_a + delta_context(a,c)", "Task context lowers the threshold for expected actions and raises it for unexpected/risky ones."],
            ["FC = sum_t cost(a_t) * I(false_t) / sum_t cost(a_t) * I(action_t)", "False action cost rate; confirmation and zoom errors carry more cost than hover."],
            ["TaskSuccess = completed_trials / all_trials", "Scenario-level result across scripted AR tasks, not only frame-level recognition."],
        ],
        [6.1, 9.9],
        header_fill=GREEN_FILL,
        font_size=8,
    )
    formula_table = doc.tables[-1]
    for row in formula_table.rows[1:]:
        for run in row.cells[0].paragraphs[0].runs:
            set_run_font(run, name="Cambria Math", size=8)


def add_section_current_state(doc: Document, data: dict, charts: dict[str, Path], screenshots: dict[str, Path]) -> None:
    doc.add_heading("2. Current State of the System", level=1)
    live = data["live"]["session"]
    official = data["official_summary"]
    add_table(
        doc,
        ["Subsystem", "Current status", "Measurement or evidence"],
        [
            [
                "Recognition",
                "C6 Ensemble is used as the robust recognizer in the current interface.",
                "Clean accuracy 0.930; clean macro F1 0.887; robust macro F1 0.859.",
            ],
            [
                "Interaction policy",
                "TARC is the proposed task-aware risk-calibrated controller.",
                f"False action cost {fmt(official['proposed_tarc']['false_action_cost_rate_mean'])}; task success {fmt(official['proposed_tarc']['task_success_rate_mean'])}.",
            ],
            [
                "Webcam prototype",
                "The local demo can run in Camera Stream mode through the webcam.",
                f"Existing live sample: {fmt(live['fps']['mean'], 2)} mean FPS, {fmt(live['processing_ms']['p95'], 2)} ms processing p95, {fmt(live['detection_rate_mean'])} detection rate.",
            ],
        ],
        [3.6, 6.2, 6.2],
        header_fill=HEADER_FILL,
        font_size=8,
    )
    c6_rows = data["c6_rows"]
    add_table(
        doc,
        ["Method", "Clean Acc.", "Clean Macro F1", "Robust Macro F1", "Robust False Action"],
        [[r["method"], r["clean_accuracy"], r["clean_macro_f1"], r["robust_macro_f1_mean"], r["robust_false_action_rate"]] for r in c6_rows],
        [5.0, 2.5, 3.0, 3.0, 2.5],
        header_fill=HEADER_FILL,
        font_size=8,
    )
    add_figure(doc, charts["macro_f1"], "Figure 1. Recognition macro F1 improved after the C6 ensemble upgrade.", width_in=6.0)

    summary = data["official_summary"]
    interaction_rows = [
        [
            "M1 Baseline Direct",
            fmt(summary["baseline_direct"]["task_success_rate_mean"]),
            fmt(summary["baseline_direct"]["action_precision_mean"]),
            fmt(summary["baseline_direct"]["unintended_action_rate_mean"]),
            fmt(summary["baseline_direct"]["false_action_cost_rate_mean"]),
        ],
        [
            "M2 Robust C6 Direct",
            fmt(summary["robust_recognizer_direct"]["task_success_rate_mean"]),
            fmt(summary["robust_recognizer_direct"]["action_precision_mean"]),
            fmt(summary["robust_recognizer_direct"]["unintended_action_rate_mean"]),
            fmt(summary["robust_recognizer_direct"]["false_action_cost_rate_mean"]),
        ],
        [
            "M3 Proposed TARC",
            fmt(summary["proposed_tarc"]["task_success_rate_mean"]),
            fmt(summary["proposed_tarc"]["action_precision_mean"]),
            fmt(summary["proposed_tarc"]["unintended_action_rate_mean"]),
            fmt(summary["proposed_tarc"]["false_action_cost_rate_mean"]),
        ],
    ]
    add_table(
        doc,
        ["Official method", "Task success", "Precision", "Unintended", "False cost"],
        interaction_rows,
        [4.8, 2.7, 2.7, 2.7, 2.7],
        header_fill=BLUE_FILL,
        font_size=8,
    )
    add_callout(
        doc,
        "Main measured result",
        "In the official replay benchmark, M3 Proposed TARC reduces false action cost from 0.110 to 0.025 while keeping "
        "task success at 0.531, approximately the same completion level as the direct baseline.",
        fill=GREEN_FILL,
    )
    add_figure(doc, charts["tradeoff"], "Figure 2. Official-method safety/completion trade-off.", width_in=6.0)
    add_figure(doc, screenshots["live"], "Figure 3. Current live interface viewport with Robust C6 and TARC selected.", width_in=6.15)
    add_figure(doc, screenshots["tables"], "Figure 4. Current experiment-results screen expanded as a wide viewport capture.", width_in=6.15)

    doc.add_heading("Places for your own live interaction screenshots", level=2)
    add_placeholder(
        doc,
        1,
        "Insert a webcam screenshot after pressing Start Task: camera stream visible, AR overlay active, Live Status showing FPS, detection and current action.",
    )
    add_placeholder(
        doc,
        2,
        "Insert a second screenshot from a completed or attempted scenario: task step state, gesture/action telemetry and scenario progress should be visible.",
    )


def add_section_future(doc: Document, data: dict) -> None:
    doc.add_heading("3. Further Development and Final Target", level=1)
    domain = data["domain"]
    doc.add_paragraph(
        "The project should finish as a documented AR gesture-control prototype with repeatable recognition metrics, "
        "task-level interaction-safety evaluation and direct visual proof of webcam-based interaction."
    )
    add_table(
        doc,
        ["Direction", "Current state", "Expected final result"],
        [
            ["Recognizer", "C6 Ensemble is the default robust recognizer in the interface.", "Use it as the main recognizer and report any live-session deviations."],
            ["Controller", "TARC is the proposed task-aware risk-calibrated controller.", "Tune latency and per-task thresholds, especially for weak scenarios."],
            ["Evaluation", "Recognition and official replay benchmark are reproducible.", "Add repeated live webcam sessions and compare logs with replay results."],
            ["Phone transfer", f"{domain['local_phone']['planned_records']} local phone-rear clips planned; {domain['local_phone']['missing_raw_video_count']} missing.", "Record local clips and rerun landmark extraction/domain-transfer metrics."],
            ["Thesis result", "Novelty is interaction-risk modeling over guided AR tasks.", "Present a cautious final claim: safer guided AR gesture control, not unrestricted open-world control."],
        ],
        [3.4, 6.4, 6.2],
        header_fill=HEADER_FILL,
        font_size=8,
    )
    add_callout(
        doc,
        "Expected final form",
        "A webcam-demonstrated and phone-transfer-tested AR gesture prototype with C6 recognition, TARC action safety, "
        "scenario-level benchmarks, live screenshots, and a thesis narrative centered on risk-aware interaction design.",
        fill=BLUE_FILL,
    )
    doc.add_heading("Final summary", level=2)
    doc.add_paragraph(
        "The current project state is academically defensible as an applied AR interaction study. The strongest originality "
        "is the task-aware risk controller and its evaluation with weighted false-action cost. C6 strengthens the recognition "
        "base, but the thesis-level value comes from showing that recognition confidence, action cost and task context can "
        "be combined into a safer AR control loop. The remaining work is empirical: insert live webcam evidence, record "
        "local phone-rear data, and validate the controller on repeated real interactions."
    )


def build_report() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    c6_rows = read_csv("artifacts/reports/c6_tables/summary.csv")
    official = read_json("artifacts/reports/official_method_benchmark.json")
    data = {
        "c6_rows": c6_rows,
        "official_summary": official["evaluation"]["summary"],
        "domain": read_json("artifacts/reports/domain_readiness.json"),
        "live": read_json("artifacts/reports/live_task_report.json"),
    }
    charts = {
        "macro_f1": create_bar_chart(
            "c6_macro_f1_chart.png",
            "Recognition clean macro F1",
            [(r["method"], float(r["clean_macro_f1"])) for r in c6_rows],
            color="#2E74B5",
            max_value=1.0,
        ),
        "tradeoff": create_tradeoff_chart(data["official_summary"]),
    }
    screenshots = {
        "live": crop_screenshot(
            ROOT / "artifacts" / "screenshots" / "report_2026_05_28_live_overview_viewport.png",
            "report_2026_05_28_live_overview_viewport.png",
        ),
        "tables": crop_screenshot(
            ROOT / "artifacts" / "screenshots" / "report_2026_05_28_tables_viewport.png",
            "report_2026_05_28_tables_viewport.png",
        ),
    }

    doc = Document()
    style_document(doc)
    add_title(doc)
    add_section_work_description(doc)
    add_section_current_state(doc, data, charts, screenshots)
    doc.add_page_break()
    add_section_future(doc, data)
    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build_report())
