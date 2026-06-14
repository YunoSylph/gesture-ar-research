from __future__ import annotations

import csv
import json
import math
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "generated"
FIG_DIR = OUT_DIR / "project_state_figures"
DOCX_PATH = OUT_DIR / "Yuzhakov_MV_Project_State_Report_2026-05-26_EN.docx"

ACCENT = RGBColor(31, 78, 121)
MUTED = RGBColor(95, 95, 95)
LIGHT_BLUE = "E8EEF5"
LIGHT_GREEN = "E2F0D9"
LIGHT_GOLD = "FFF2CC"
LIGHT_RED = "FCE4D6"
LIGHT_GRAY = "F2F4F7"


def read_json(relative: str) -> dict:
    return json.loads((ROOT / relative).read_text(encoding="utf-8"))


def read_csv(relative: str) -> list[dict[str, str]]:
    with (ROOT / relative).open(encoding="utf-8", newline="") as fh:
        return list(csv.DictReader(fh))


def fmt(value: float | str | int | None, digits: int = 4) -> str:
    if value is None or value == "":
        return "-"
    if isinstance(value, str):
        try:
            value = float(value)
        except ValueError:
            return value
    return f"{float(value):.{digits}f}"


def fmt_pct(value: float | str | None, digits: int = 1) -> str:
    if value is None or value == "":
        return "-"
    if isinstance(value, str):
        value = float(value)
    return f"{float(value) * 100:.{digits}f}%"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        tag = "w:" + edge
        element = tc_mar.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold: bool = False, size: int = 9, center: bool | None = None) -> None:
    cell.text = ""
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if center is None:
        center = len(text) < 28
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color: str = "C8CED8") -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_width(table, widths_cm: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_cm: list[float],
    header_fill: str = LIGHT_BLUE,
    font_size: int = 9,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    set_row_cant_split(hdr)
    for idx, header in enumerate(headers):
        set_cell_text(hdr.cells[idx], header, bold=True, size=font_size)
        set_cell_shading(hdr.cells[idx], header_fill)
    for row in rows:
        row_obj = table.add_row()
        set_row_cant_split(row_obj)
        cells = row_obj.cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value), size=font_size)
    set_table_width(table, widths_cm)
    doc.add_paragraph()


def add_formula_table(doc: Document, rows: list[list[str]]) -> None:
    add_table(
        doc,
        ["No.", "Formula", "Meaning in this project"],
        rows,
        widths_cm=[1.2, 7.6, 7.0],
        header_fill=LIGHT_GREEN,
        font_size=9,
    )
    # Apply math-friendly font to formula cells.
    table = doc.tables[-1]
    for row in table.rows[1:]:
        for run in row.cells[1].paragraphs[0].runs:
            run.font.name = "Cambria Math"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")


def get_font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for item in candidates:
        if Path(item).exists():
            return ImageFont.truetype(item, size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = word if not current else current + " " + word
        bbox = draw.textbbox((0, 0), candidate, font=font)
        if bbox[2] - bbox[0] <= width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_box(draw, box, text, fill, outline="#4F81BD", font=None) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=12, fill=fill, outline=outline, width=2)
    font = font or get_font(24, True)
    small = get_font(18)
    lines = wrap_text(draw, text, font, x2 - x1 - 32)
    if len(lines) > 2:
        font = small
        lines = wrap_text(draw, text, font, x2 - x1 - 32)
    line_h = font.size + 5 if hasattr(font, "size") else 24
    total_h = line_h * len(lines)
    y = y1 + (y2 - y1 - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, y), line, fill="#1F1F1F", font=font)
        y += line_h


def draw_arrow(draw, start, end, color="#4F81BD") -> None:
    draw.line([start, end], fill=color, width=4)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 14
    p1 = (end[0] - size * math.cos(angle - math.pi / 6), end[1] - size * math.sin(angle - math.pi / 6))
    p2 = (end[0] - size * math.cos(angle + math.pi / 6), end[1] - size * math.sin(angle + math.pi / 6))
    draw.polygon([end, p1, p2], fill=color)


def make_architecture_figure(path: Path) -> None:
    img = Image.new("RGB", (1600, 920), "white")
    draw = ImageDraw.Draw(img)
    title = get_font(40, True)
    draw.text((60, 42), "Current Project Architecture: Public Data, Hybrid Recognition, AR Interaction", fill="#1F4E79", font=title)
    boxes = [
        ((70, 150, 345, 270), "IPN Hand + planned phone rear AR clips", "#E8EEF5"),
        ((445, 150, 720, 270), "MediaPipe hand landmarks [T,21,3]", "#E2F0D9"),
        ((820, 150, 1095, 270), "Dual-view features Z ∈ R^(32×74)", "#FFF2CC"),
        ((1195, 150, 1515, 270), "C1-T Temporal TCN", "#FCE4D6"),
        ((445, 420, 720, 540), "C3 geometry-aware safety prior", "#E2F0D9"),
        ((820, 420, 1095, 540), "C2 context gate", "#E8EEF5"),
        ((1195, 420, 1515, 540), "React + Three.js AR interface", "#FFF2CC"),
        ((820, 680, 1095, 800), "ONNX now; Core ML contract next", "#F2F4F7"),
    ]
    for box, text, fill in boxes:
        draw_box(draw, box, text, fill)
    arrows = [
        ((345, 210), (445, 210)),
        ((720, 210), (820, 210)),
        ((1095, 210), (1195, 210)),
        ((1355, 270), (1355, 420)),
        ((1195, 480), (1095, 480)),
        ((820, 480), (720, 480)),
        ((955, 540), (955, 680)),
        ((955, 270), (955, 420)),
    ]
    for start, end in arrows:
        draw_arrow(draw, start, end)
    note_font = get_font(23)
    note = (
        "The research contribution is not a standalone classifier only: it is a measured path "
        "from temporal recognition to safer AR actions under context and perturbation."
    )
    for idx, line in enumerate(wrap_text(draw, note, note_font, 1380)):
        draw.text((90, 850 + idx * 28), line, fill="#595959", font=note_font)
    img.save(path)


def make_final_target_roadmap(path: Path) -> None:
    img = Image.new("RGB", (1600, 820), "white")
    draw = ImageDraw.Draw(img)
    title = get_font(40, True)
    draw.text((60, 42), "Development Roadmap Toward the Final System", fill="#1F4E79", font=title)
    stages = [
        ("Current baseline", "Validated C1-T + ONNX + AR UI", "#E8EEF5"),
        ("C3 research hardening", "Robustness, ablation, calibration, CI tables", "#E2F0D9"),
        ("Phone AR domain", "35 rear-camera clips + zero-shot evaluation", "#FFF2CC"),
        ("User study", "Direct vs C2 vs C3+C2 task metrics", "#FCE4D6"),
        ("Final deliverable", "Mobile-ready AR gesture interaction thesis system", "#D9EAD3"),
    ]
    x = 65
    y = 210
    w = 275
    h = 170
    for idx, (head, body, fill) in enumerate(stages):
        bx = x + idx * 300
        draw.rounded_rectangle((bx, y, bx + w, y + h), radius=16, fill=fill, outline="#4F81BD", width=2)
        draw.text((bx + 24, y + 24), head, fill="#1F1F1F", font=get_font(25, True))
        for li, line in enumerate(wrap_text(draw, body, get_font(20), w - 48)):
            draw.text((bx + 24, y + 72 + li * 26), line, fill="#404040", font=get_font(20))
        if idx < len(stages) - 1:
            draw_arrow(draw, (bx + w, y + h // 2), (bx + 300, y + h // 2))
    lower = [
        ("Research proof", "Bootstrap confidence intervals and scenario-specific robustness deltas."),
        ("Engineering proof", "Stable camera loop, exported model contracts, reproducible reports."),
        ("Interaction proof", "Task success, unintended-action rate, false triggers per minute, user workload."),
    ]
    for idx, (head, body) in enumerate(lower):
        bx = 115 + idx * 485
        draw.rounded_rectangle((bx, 520, bx + 405, 660), radius=12, fill="#F7FBFE", outline="#9ECAE1", width=2)
        draw.text((bx + 22, 545), head, fill="#1F4E79", font=get_font(24, True))
        for li, line in enumerate(wrap_text(draw, body, get_font(18), 360)):
            draw.text((bx + 22, 588 + li * 24), line, fill="#404040", font=get_font(18))
    img.save(path)


def normalize_image(src: Path, dst: Path, max_width: int = 1600) -> Path:
    dst.parent.mkdir(parents=True, exist_ok=True)
    with Image.open(src) as im:
        im = im.convert("RGB")
        if im.width > max_width:
            ratio = max_width / im.width
            im = im.resize((max_width, int(im.height * ratio)))
        im.save(dst, quality=92)
    return dst


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.color.rgb = ACCENT if name != "Heading 3" else RGBColor(31, 77, 120)
        style.font.bold = True
    doc.styles["Heading 1"].font.size = Pt(16)
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(16)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(8)
    doc.styles["Heading 2"].font.size = Pt(13)
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(12)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(6)
    doc.styles["Heading 3"].font.size = Pt(12)
    doc.styles["Heading 3"].paragraph_format.space_before = Pt(8)
    doc.styles["Heading 3"].paragraph_format.space_after = Pt(4)

    header = section.header.paragraphs[0]
    header.text = "Gesture AR Project State Report"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(11)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.10
        run = p.add_run(item)
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        run.font.size = Pt(11)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_picture(doc: Document, path: Path, width: float, caption: str) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Gesture AR Project State Report")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("Work description, current system state, and target development path")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(13)
    run.font.color.rgb = MUTED

    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Project", "Context-Aware Temporal Landmark Gesture Recognition for AR Interaction"],
            ["Prepared for", "Current research and thesis planning snapshot"],
            ["Prepared on", "26 May 2026"],
            ["Language", "English"],
            ["Design preset", "standard_business_brief"],
        ],
        widths_cm=[4.0, 11.8],
        header_fill=LIGHT_GRAY,
        font_size=10,
    )
    add_para(
        doc,
        "This document summarizes the current project state as a thesis-oriented engineering and research report. "
        "It is not a final dissertation chapter set; it is a structured snapshot used to describe the work, the validated system state, and the intended final target."
    )
    doc.add_page_break()


def add_contents(doc: Document) -> None:
    add_heading(doc, "Contents", 1)
    rows = [
        ["1", "Work Description"],
        ["2", "Current System State"],
        ["3", "Further Development and Final Target"],
    ]
    add_table(doc, ["Section", "Title"], rows, widths_cm=[2.0, 13.8], header_fill=LIGHT_GRAY, font_size=10)
    doc.add_page_break()


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)

    valid = read_json("artifacts/reports/ipn_c1t_tcn_full_validated_recognition.json")
    risk = read_json("artifacts/reports/recognition_risk_analysis.json")
    domain = read_json("artifacts/reports/domain_readiness.json")
    live = read_json("artifacts/reports/live_task_report.json")
    preprocess = read_json("artifacts/mobile/gesture_mobile_bundle/preprocessing_contract.json")
    c2 = read_json("artifacts/mobile/gesture_mobile_bundle/c2_policy.json")
    portability = read_json("artifacts/mobile/gesture_mobile_bundle/portability_contract.json")
    robustness = read_csv("artifacts/reports/c3_tables/c3_robustness_summary.csv")
    ablation = read_csv("artifacts/reports/c3_tables/c3_ablation_summary.csv")
    policy = read_csv("artifacts/reports/c3_tables/c3_policy_ablation.csv")
    calibration = read_csv("artifacts/reports/c3_tables/c3_calibration_candidates.csv")

    arch = FIG_DIR / "project_architecture_en.png"
    roadmap = FIG_DIR / "final_target_roadmap_en.png"
    make_architecture_figure(arch)
    make_final_target_roadmap(roadmap)

    figure_paths = {
        "robustness": ROOT / "artifacts/figures/c3_robustness_macro_f1.png",
        "ablation": ROOT / "artifacts/figures/c3_ablation_perturbed_macro_f1.png",
        "policy": ROOT / "artifacts/figures/c3_policy_unintended_action_rate.png",
        "results": ROOT / "artifacts/screenshots/c3_ablation_results_page.png",
        "sorting": ROOT / "artifacts/screenshots/ar_interface_virtual_sorting_final_checked.png",
    }
    normalized = {
        key: normalize_image(path, FIG_DIR / f"{key}.jpg", max_width=1500)
        for key, path in figure_paths.items()
        if path.exists()
    }

    doc = Document()
    configure_styles(doc)
    add_title_page(doc)
    add_contents(doc)

    add_heading(doc, "1. Work Description", 1)
    add_para(
        doc,
        "The project develops a Windows-first research pipeline for context-aware gesture recognition in augmented reality. "
        "The core idea is to recognize hand gestures from temporal sequences of 21 hand landmarks, transform recognition outputs into safer AR actions, "
        "and measure the result not only by classifier accuracy but also by interaction-level risk."
    )
    add_para(
        doc,
        "The work is structured around three layers: a temporal recognition layer, a hybrid safety layer, and an interaction policy layer. "
        "The temporal layer uses a compact TCN over normalized landmark features. The C3 Hybrid layer adds a lightweight geometry-aware prior and a safety gate. "
        "The C2 layer converts gesture predictions into AR actions only after confidence, stability, and cooldown constraints are satisfied."
    )
    add_picture(doc, arch, 6.5, "Figure 1. Project architecture and data-to-action flow.")
    add_table(
        doc,
        ["Component", "Current role", "Why it matters"],
        [
            ["IPN Hand data", "Public benchmark for reproducible recognition and robustness experiments", "Avoids relying only on a private demonstration dataset"],
            ["Landmark tensor", "MediaPipe-style [T,21,3] representation", "Reduces visual complexity and enables fast temporal modeling"],
            ["C1-T TCN", "Temporal classifier over [32,74] features", "Main learned recognizer and ONNX-exported runtime path"],
            ["C3 Hybrid", "Temporal probabilities plus geometry-aware safety prior", "Targets false-action risk rather than only clean accuracy"],
            ["C2 Gate", "Context-aware gesture-to-action finite-state policy", "Suppresses unstable or repeated accidental commands"],
            ["AR UI", "React + Three.js task interface", "Turns recognition into measurable AR task execution"],
        ],
        widths_cm=[3.4, 6.2, 6.2],
    )
    add_para(
        doc,
        "The gesture vocabulary contains seven target classes: no_gesture, point_2f, click_2f, swipe_left, swipe_right, zoom_in, and zoom_out. "
        "These classes are deliberately small but operationally complete for pointing, selection, navigation, scaling, and suppression of accidental input."
    )
    add_formula_table(
        doc,
        [
            ["(1)", "P(t) = [p(t,1), ..., p(t,21)],  p(t,j)=(x,y,z);  X ∈ R^(T×21×3)", "Each frame is represented as 21 hand landmarks, and each gesture clip is a temporal landmark tensor."],
            ["(2)", "s(t)=max(||p5_xy-p17_xy||_2, ||p9_xy-p0_xy||_2, ε)", "Palm scale normalizes the hand pose and guards against division by zero."],
            ["(3)", "p_norm(t,j)=(p(t,j)-p(t,0))/s(t)", "Landmarks are represented relative to the wrist and palm scale."],
            ["(4)", "z(t)=[vec(P_norm), motion, confidence] ∈ R^74;  Z ∈ R^(32×74)", "The model input combines 63 pose features and 11 motion/confidence features."],
        ],
    )

    add_heading(doc, "2. Current System State", 1)
    add_para(
        doc,
        "As of 26 May 2026, the project has moved beyond a simple prototype. It contains extracted full IPN landmark benchmarks, trained C1 and C1-T models, "
        "validated TCN recognition reports, ONNX export, mobile portability contracts, AR task scenarios, live-session reports, C3 Hybrid robustness experiments, "
        "ablation tables, calibration candidates, charts, and UI screenshots."
    )
    rec = valid["recognition"]
    lat = valid["latency"]
    c3_row = next(row for row in robustness if row["method"] == "c3_hybrid")
    c1t_row = next(row for row in robustness if row["method"] == "c1t_direct")
    best_calib = calibration[0]
    add_table(
        doc,
        ["Measurement", "Value", "Interpretation"],
        [
            ["Full IPN subset", "2405 train / 1033 test clips", "Primary reproducible public benchmark"],
            ["Validated C1-T accuracy", fmt(rec["accuracy"]), "Recognition quality on the full test subset"],
            ["Validated C1-T macro F1", fmt(rec["macro_f1"]), "Class-balanced recognition quality"],
            ["Validated C1-T p95 latency", f"{lat['offline_latency_ms_p95']:.3f} ms", "Offline per-sample inference latency"],
            ["C3 clean macro F1", fmt(c3_row["clean_macro_f1"]), "Hybrid clean benchmark score"],
            ["C3 perturbed macro F1 mean", fmt(c3_row["perturbed_macro_f1_mean"]), "Average robustness under synthetic perturbations"],
            ["C3 mean perturbed false-action rate", fmt(c3_row["perturbed_false_action_rate_mean"]), "Interaction-relevant risk signal"],
        ],
        widths_cm=[4.5, 4.7, 6.6],
    )
    add_para(
        doc,
        f"The current best C3 calibration candidate uses neural_weight={best_calib['neural_weight']}, "
        f"geometry_weight={best_calib['geometry_weight']}, and action_threshold={best_calib['action_threshold']}."
    )
    add_picture(doc, normalized["results"], 6.5, "Figure 2. Current experiment results page with C3 Hybrid state.")
    add_para(
        doc,
        "The current C3 result should be interpreted as a robustness and interaction-risk improvement rather than a dramatic clean-accuracy jump. "
        "The validated temporal recognizer already performs strongly on clean public data; the C3 contribution is to make action emission more conservative where AR interaction risk matters."
    )
    add_table(
        doc,
        ["Method", "Clean macro F1", "Perturbed macro F1 mean", "Macro F1 drop", "Perturbed false-action rate"],
        [
            [row["method"], fmt(row["clean_macro_f1"]), fmt(row["perturbed_macro_f1_mean"]), fmt(row["macro_f1_drop"]), fmt(row["perturbed_false_action_rate_mean"])]
            for row in robustness
        ],
        widths_cm=[3.4, 3.0, 3.7, 2.8, 3.0],
    )
    add_picture(doc, normalized["robustness"], 6.5, "Figure 3. C3 robustness by perturbation scenario.")
    add_table(
        doc,
        ["Method", "Clean macro F1", "Perturbed macro F1 mean", "Macro F1 drop", "Perturbed false-action rate"],
        [
            [row["method"], fmt(row["clean_macro_f1"]), fmt(row["perturbed_macro_f1_mean"]), fmt(row["macro_f1_drop"]), fmt(row["perturbed_false_action_rate_mean"])]
            for row in ablation
        ],
        widths_cm=[3.6, 3.0, 3.7, 2.8, 2.8],
    )
    add_picture(doc, normalized["ablation"], 6.5, "Figure 4. Recognition ablation under perturbations.")
    add_table(
        doc,
        ["Method", "Policy", "Action precision", "Action recall", "Unintended action rate", "False triggers / min"],
        [
            [
                row["method"],
                row["policy"],
                fmt(row["action_precision"]),
                fmt(row["action_recall"]),
                fmt(row["unintended_action_rate"]),
                fmt(row["false_trigger_rate_per_minute"], digits=3),
            ]
            for row in policy
        ],
        widths_cm=[2.8, 2.1, 2.7, 2.5, 3.2, 2.5],
    )
    add_picture(doc, normalized["policy"], 6.5, "Figure 5. Interaction-policy false action risk.")
    live_session = live["session"]
    placement = live["tasks"]["placement"]
    add_table(
        doc,
        ["Live/task metric", "Value", "Meaning"],
        [
            ["Session", live_session["session_id"], "Latest task-level live report used as a diagnostic snapshot"],
            ["Frames / duration", f"{live_session['frames']} / {live_session['duration_ms'] / 1000:.2f} s", "Captured live loop length"],
            ["FPS mean / p95", f"{live_session['fps']['mean']:.2f} / {live_session['fps']['p95']:.2f}", "Camera plus backend loop rate"],
            ["Processing mean / p95", f"{live_session['processing_ms']['mean']:.3f} / {live_session['processing_ms']['p95']:.3f} ms", "Backend processing time"],
            ["Detection-rate mean", fmt(live_session["detection_rate_mean"]), "Average hand-detection availability"],
            ["Pointer coverage", fmt(placement["pointer_coverage"]), "Frames where AR pointer coordinates were available"],
            ["Required-action coverage", fmt(placement["required_action_coverage"]), "Scenario completion was not yet achieved"],
        ],
        widths_cm=[4.4, 4.4, 7.0],
    )
    add_picture(doc, normalized["sorting"], 4.3, "Figure 6. Virtual Sorting AR scenario surface.")
    add_formula_table(
        doc,
        [
            ["(5)", "p_k = softmax(W · GAP(TCN(Z)) + b)_k;  y_hat = argmax_k p_k", "Temporal recognition maps the [32,74] sequence to a gesture label."],
            ["(6)", "q_k = softmax(w_n log p_k + w_g r_k)", "C3 Hybrid fuses neural probabilities with geometry-aware prior scores."],
            ["(7)", "a(t)=g(y_hat(t)) if p_max≥τ, stable≥S, and t-t_last≥Δ; otherwise a(t)=∅", "C2 emits an AR action only after confidence, stability, and cooldown checks."],
            ["(8)", "FAR_no = FP_action(no_gesture) / N_no_gesture", "False-action risk is measured specifically for frames/clips with no intended gesture."],
            ["(9)", "Coverage_req = |A_req ∩ A_obs| / |A_req|", "Task-level coverage measures whether the expected action set appeared in the live scenario."],
        ],
    )
    add_table(
        doc,
        ["Portability item", "Current state"],
        [
            ["Model input contract", str(preprocess["model_input_shape"])],
            ["Pose representation", preprocess["pose_features"]["normalization"]],
            ["Motion feature count", str(preprocess["motion_features"]["dim"])],
            ["C2 threshold / stable frames / cooldown", f"{c2['activation_threshold']} / {c2['stable_frames']} / {c2['cooldown_ms']} ms"],
            ["ONNX export", "Available"],
            ["Core ML stage", portability["coreml_stage"]["status"]],
            ["Phone domain status", domain["domain_transfer_status"]],
            ["Planned phone rear AR clips", str(domain["local_phone"]["planned_records"])],
            ["Missing phone raw videos", str(domain["local_phone"]["missing_raw_video_count"])],
        ],
        widths_cm=[5.2, 10.6],
    )

    add_heading(doc, "3. Further Development and Final Target", 1)
    add_para(
        doc,
        "The final project should be a defensible master-level system rather than only a demonstration. "
        "It should show that temporal landmark recognition, a geometry-aware safety prior, and a context-aware action policy jointly improve practical AR interaction."
    )
    add_picture(doc, roadmap, 6.5, "Figure 7. Roadmap from current state to final thesis-ready system.")
    add_table(
        doc,
        ["Development area", "Next work", "Expected final evidence"],
        [
            ["Phone rear AR data", "Record the planned 35 local clips and extract landmarks", "Zero-shot and calibrated domain-shift report"],
            ["C3 validation", "Add bootstrap confidence intervals and scenario-specific sensitivity analysis", "Statistically interpretable robustness section"],
            ["C2 policy", "Tune threshold, stable frames, and cooldown against task metrics", "Direct vs C2 vs C3+C2 interaction comparison"],
            ["AR task benchmark", "Expand replay/live scenarios beyond placement and Virtual Sorting", "Task success, false triggers/min, corrections/task"],
            ["Mobile runtime", "Complete Core ML conversion and iOS/RealityKit integration", "Phone AR prototype using the same preprocessing contract"],
            ["Thesis package", "Freeze scripts, reports, charts, and reproducible commands", "Repeatable experimental chapter and final defense artifacts"],
        ],
        widths_cm=[4.0, 6.4, 5.4],
    )
    add_formula_table(
        doc,
        [
            ["(10)", "Accuracy = trace(M) / sum(M)", "Base recognition correctness on public and local domains."],
            ["(11)", "MacroF1=(1/K)·Σ_i F1(i);  WeightedF1=Σ_i n_i·F1(i)/Σ_i n_i", "Balanced and support-weighted quality for imbalanced gesture classes."],
            ["(12)", "UAR = N_unintended / N_actions", "Unintended-action rate is the main AR safety metric for interaction policy comparison."],
            ["(13)", "Gap_D(m)=|m_IPN - m_phone|", "Domain gap quantifies transfer from public IPN Hand to phone rear-camera AR clips."],
            ["(14)", "Score = m_robust - λ·FAR - μ·Latency", "Final model selection should balance robustness, false actions, and runtime."],
        ],
    )
    add_para(
        doc,
        "The intended final system is an AR gesture interaction stack that can be described, reproduced, measured, and defended: "
        "public-data training and robustness, local phone-domain validation, ONNX/Core ML portability, task-level AR evaluation, and a clear statement of remaining limitations."
    )

    for section in doc.sections:
        section.start_type = WD_SECTION_START.NEW_PAGE if section is not doc.sections[0] else section.start_type

    doc.core_properties.title = "Gesture AR Project State Report"
    doc.core_properties.subject = "Current project state, measurements, formulas, charts, and development plan"
    doc.core_properties.author = "Maksim Iuzhakov"
    doc.core_properties.keywords = "AR, gesture recognition, TCN, C3 Hybrid, C2 Gate, landmarks, ONNX, Core ML"
    doc.core_properties.created = datetime(2026, 5, 26, 0, 0, 0)
    doc.save(DOCX_PATH)
    with zipfile.ZipFile(DOCX_PATH) as zf:
        bad = zf.testzip()
        if bad:
            raise RuntimeError(f"Invalid DOCX zip entry: {bad}")
    print(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
